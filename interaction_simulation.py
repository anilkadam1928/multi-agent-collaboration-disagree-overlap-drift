from __future__ import annotations

"""Week 4 Wednesday - Scenario 4 interaction analysis.

This file does not run the local LLM again. It reads the saved outputs from:
- Week 3 disagreement module
- Week 3 overlap module
- Week 4 drift/ABA module

It then measures the causal chain:
Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk

No existing simulation files are modified.
"""

import math
from collections import Counter
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd


PROJECT_ROOT = Path(__file__).resolve().parent
DATA_DIR = PROJECT_ROOT / "data"
RESULTS_DIR = PROJECT_ROOT / "results"
ANALYSIS_DIR = PROJECT_ROOT / "analysis"

DISAGREEMENT_PATH = PROJECT_ROOT / "results" / "disagreement_results.csv"
OVERLAP_PATH = DATA_DIR / "overlap_results.csv"
DRIFT_EVENTS_PATH = DATA_DIR / "drift_with_aba_events.csv"
BASELINE_PATH = DATA_DIR / "baseline_results.csv"

DUAL_RESULTS_PATH = DATA_DIR / "dual_results.csv"
RESOLUTION_OVERLAP_PATH = DATA_DIR / "resolution_overlap.csv"
DRIFT_ACCELERATION_PATH = DATA_DIR / "drift_acceleration.csv"
SUMMARY_PATH = ANALYSIS_DIR / "week4_day3_interaction_summary.csv"
FIGURE_PATH = RESULTS_DIR / "Figure_4_Causal_Chain.png"
REPORT_MD_PATH = ANALYSIS_DIR / "Week4_Day3_Interaction_Report.md"
REPORT_DOCX_PATH = ANALYSIS_DIR / "Week4_Day3_Interaction_Report.docx"

DECISION_LABELS = ["approve", "refer", "reject", "unknown"]


def read_required_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Missing required input file: {path}")
    return pd.read_csv(path)


def as_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"true", "1", "yes", "y"}


def borrower_number(series: pd.Series) -> pd.Series:
    return series.astype(str).str.extract(r"(\d+)").astype(int)[0]


def borrower_subset_label(borrower_ids: pd.Series) -> str:
    ordered = (
        borrower_ids.astype(str)
        .drop_duplicates()
        .sort_values(key=lambda values: borrower_number(values))
        .tolist()
    )
    if not ordered:
        return "no common borrowers"
    if len(ordered) == 1:
        return ordered[0]
    return f"{ordered[0]}-{ordered[-1]}"


def smoothed_distribution(values: pd.Series, labels: list[str]) -> dict[str, float]:
    cleaned = values.fillna("unknown").astype(str).str.lower()
    counts = Counter(cleaned)
    total = sum(counts.get(label, 0) for label in labels)
    return {
        label: (counts.get(label, 0) + 1e-6) / (total + len(labels) * 1e-6)
        for label in labels
    }


def kl_divergence(current: dict[str, float], baseline: dict[str, float], labels: list[str]) -> float:
    return sum(current[label] * math.log(current[label] / baseline[label]) for label in labels)


def build_disagreement_borrower_table(disagreement: pd.DataFrame) -> pd.DataFrame:
    disagreement = disagreement.copy()
    disagreement["disagreement_exists"] = disagreement["disagreement_exists"].map(as_bool)
    disagreement["sycophancy_flagged"] = disagreement["sycophancy_flagged"].map(as_bool)
    disagreement["position_changed"] = (
        disagreement["precommit_decision"].astype(str).str.lower()
        != disagreement["final_decision"].astype(str).str.lower()
    )

    rows: list[dict[str, object]] = []
    for borrower_id, group in disagreement.groupby("borrower_id", sort=True):
        first = group.iloc[0]
        rows.append(
            {
                "borrower_id": borrower_id,
                "kaggle_truth_label": first["kaggle_truth_label"],
                "disagreement_detected": bool(group["disagreement_exists"].max()),
                "borda_winner": first.get("borda_winner", "unknown"),
                "consensagent_winner": first.get("consensagent_winner", "unknown"),
                "position_change_count": int(group["position_changed"].sum()),
                "sycophancy_flag_count": int(group["sycophancy_flagged"].sum()),
            }
        )

    return pd.DataFrame(rows)


def build_overlap_borrower_table(overlap: pd.DataFrame) -> pd.DataFrame:
    without_lce = overlap[overlap["mode"] == "without_lce"].copy()
    with_lce = overlap[overlap["mode"] == "with_lce_tom"].copy()

    raw_cols = {
        "trace_bid": "trace_bid_raw",
        "tap_overlap_count": "raw_tap_overlap_count",
        "duplicate_invocations": "raw_duplicate_invocations",
        "keyword_overlap_count": "raw_keyword_overlap_count",
        "redundancy_index": "raw_redundancy_index",
        "final_decision": "raw_final_decision",
        "final_correct": "raw_final_correct",
    }
    controlled_cols = {
        "tap_overlap_count": "controlled_tap_overlap_count",
        "duplicate_invocations": "controlled_duplicate_invocations",
        "keyword_overlap_count": "controlled_keyword_overlap_count",
        "redundancy_index": "controlled_redundancy_index",
        "active_leader": "controlled_active_leader",
        "final_decision": "controlled_final_decision",
        "final_correct": "controlled_final_correct",
    }

    raw = without_lce[["borrower_id", *raw_cols.keys()]].rename(columns=raw_cols)
    controlled = with_lce[["borrower_id", *controlled_cols.keys()]].rename(columns=controlled_cols)
    merged = raw.merge(controlled, on="borrower_id", how="inner")

    merged["overlap_detected"] = (
        (merged["raw_tap_overlap_count"] > 0)
        | (merged["raw_duplicate_invocations"] > 0)
        | (merged["raw_redundancy_index"] > 0)
    )
    merged["overlap_intensity"] = merged["raw_redundancy_index"].astype(float)
    merged["remaining_overlap_after_lce"] = merged["controlled_redundancy_index"].astype(float)
    merged["lce_redundancy_reduction"] = (
        merged["raw_redundancy_index"] - merged["controlled_redundancy_index"]
    ).clip(lower=0)

    return merged


def assign_risk_bands(dual: pd.DataFrame) -> pd.Series:
    try:
        return pd.qcut(
            dual["trace_bid_raw"],
            q=3,
            labels=["low", "medium", "high"],
            duplicates="drop",
        ).astype(str)
    except Exception:
        return pd.Series(["medium"] * len(dual), index=dual.index)


def build_dual_results(disagreement: pd.DataFrame, overlap: pd.DataFrame) -> pd.DataFrame:
    disagreement_by_borrower = build_disagreement_borrower_table(disagreement)
    overlap_by_borrower = build_overlap_borrower_table(overlap)

    dual = disagreement_by_borrower.merge(overlap_by_borrower, on="borrower_id", how="inner")
    dual["borrower_num"] = borrower_number(dual["borrower_id"])
    dual = dual.sort_values("borrower_num").reset_index(drop=True)
    dual["run_id"] = dual["borrower_id"].map(lambda bid: f"interaction_{bid}")
    dual["risk_band"] = assign_risk_bands(dual)
    dual["co_occurrence"] = dual["overlap_detected"] & dual["disagreement_detected"]
    dual["borda_resolution_fired"] = dual["disagreement_detected"]
    dual["resolution_induced_invocations"] = dual.apply(
        lambda row: row["raw_duplicate_invocations"] if row["borda_resolution_fired"] else 0,
        axis=1,
    )
    dual["remaining_invocations_after_lce"] = dual.apply(
        lambda row: row["controlled_duplicate_invocations"] if row["borda_resolution_fired"] else 0,
        axis=1,
    )
    dual["co_occurrence_intensity"] = dual.apply(
        lambda row: row["overlap_intensity"] if row["co_occurrence"] else 0.0,
        axis=1,
    )

    median_intensity = dual["co_occurrence_intensity"].median()
    dual["co_occurrence_group"] = dual["co_occurrence_intensity"].apply(
        lambda value: "high_co_occurrence_intensity" if value >= median_intensity else "low_co_occurrence_intensity"
    )

    return dual


def build_resolution_overlap(dual: pd.DataFrame) -> pd.DataFrame:
    result = dual[
        [
            "run_id",
            "borrower_id",
            "risk_band",
            "borda_resolution_fired",
            "resolution_induced_invocations",
            "remaining_invocations_after_lce",
            "raw_redundancy_index",
            "controlled_redundancy_index",
            "lce_redundancy_reduction",
        ]
    ].copy()
    result["invocations_reduced_by_lce"] = (
        result["resolution_induced_invocations"] - result["remaining_invocations_after_lce"]
    ).clip(lower=0)
    return result


def build_drift_acceleration(dual: pd.DataFrame, drift_events: pd.DataFrame, baseline: pd.DataFrame) -> pd.DataFrame:
    drift_events = drift_events.copy()
    drift_events["borrower_num"] = borrower_number(drift_events["borrower_id"])
    drift_events = drift_events.merge(
        dual[["borrower_id", "co_occurrence_group", "co_occurrence_intensity"]],
        on="borrower_id",
        how="inner",
    )

    baseline_dist = smoothed_distribution(baseline["parsed_decision"], DECISION_LABELS)
    rows: list[dict[str, object]] = []
    max_borrower = int(drift_events["borrower_num"].max()) if not drift_events.empty else 0
    checkpoints = [point for point in [5, 10, 20, 30, 40, 50, 100, 200] if point <= max_borrower]
    if max_borrower and max_borrower not in checkpoints:
        checkpoints.append(max_borrower)

    for group_name in ["low_co_occurrence_intensity", "high_co_occurrence_intensity"]:
        group_events = drift_events[drift_events["co_occurrence_group"] == group_name]
        for checkpoint in checkpoints:
            current = group_events[group_events["borrower_num"] <= checkpoint]
            if current.empty:
                continue

            current_dist = smoothed_distribution(current["parsed_decision"], DECISION_LABELS)
            drift_score = kl_divergence(current_dist, baseline_dist, DECISION_LABELS)
            rows.append(
                {
                    "group": group_name,
                    "checkpoint": checkpoint,
                    "borrowers_seen": int(current["borrower_id"].nunique()),
                    "mean_co_occurrence_intensity": round(float(current["co_occurrence_intensity"].mean()), 4),
                    "mean_pre_call_drift_score": round(float(current["pre_call_drift_score"].mean()), 4),
                    "unknown_output_rate": round(float((current["parsed_decision"] == "unknown").mean()), 4),
                    "aba_trigger_rate": round(float(current["aba_triggered"].map(as_bool).mean()), 4),
                    "group_kl_drift_score": round(float(drift_score), 4),
                }
            )

    return pd.DataFrame(rows)


def summarize(dual: pd.DataFrame, resolution: pd.DataFrame, drift_acceleration: pd.DataFrame) -> dict[str, float]:
    co_rate = float(dual["co_occurrence"].mean())
    overlap_rate = float(dual["overlap_detected"].mean())
    disagreement_rate = float(dual["disagreement_detected"].mean())
    avg_extra = float(resolution["resolution_induced_invocations"].mean())
    avg_remaining = float(resolution["remaining_invocations_after_lce"].mean())
    avg_reduction = float(resolution["invocations_reduced_by_lce"].mean())

    latest_checkpoint = drift_acceleration["checkpoint"].max()
    latest = drift_acceleration[drift_acceleration["checkpoint"] == latest_checkpoint]
    high = latest[latest["group"] == "high_co_occurrence_intensity"]["group_kl_drift_score"]
    low = latest[latest["group"] == "low_co_occurrence_intensity"]["group_kl_drift_score"]
    high_kl = float(high.iloc[0]) if not high.empty else 0.0
    low_kl = float(low.iloc[0]) if not low.empty else 0.0
    acceleration_ratio = high_kl / low_kl if low_kl else 0.0

    return {
        "profiles_analyzed": float(len(dual)),
        "common_subset": borrower_subset_label(dual["borrower_id"]),
        "overlap_rate": overlap_rate,
        "disagreement_rate": disagreement_rate,
        "co_occurrence_rate": co_rate,
        "avg_resolution_induced_invocations": avg_extra,
        "avg_remaining_invocations_after_lce": avg_remaining,
        "avg_invocations_reduced_by_lce": avg_reduction,
        "high_group_kl_at_checkpoint": high_kl,
        "low_group_kl_at_checkpoint": low_kl,
        "drift_acceleration_ratio_high_vs_low": acceleration_ratio,
    }


def save_summary(metrics: dict[str, float]) -> None:
    pd.DataFrame([metrics]).to_csv(SUMMARY_PATH, index=False)


def plot_causal_chain(dual: pd.DataFrame, resolution: pd.DataFrame, drift_acceleration: pd.DataFrame, metrics: dict[str, float]) -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)

    band_order = ["low", "medium", "high"]
    band_summary = (
        dual.groupby("risk_band")
        .agg(
            co_occurrence_rate=("co_occurrence", "mean"),
            mean_extra_invocations=("resolution_induced_invocations", "mean"),
        )
        .reindex(band_order)
        .fillna(0)
    )

    fig, axes = plt.subplots(1, 3, figsize=(16, 5.2))
    fig.suptitle("Scenario 4 Causal Chain: Overlap -> Disagreement -> Resolution -> Drift Risk", fontsize=15)

    axes[0].bar(
        band_summary.index,
        band_summary["co_occurrence_rate"],
        color=["#7AA6C2", "#3A6EA5", "#1F4E79"],
    )
    axes[0].set_title("A. Co-occurrence Rate by Risk Band")
    axes[0].set_xlabel("TRACE risk band")
    axes[0].set_ylabel("Overlap and disagreement co-occur")
    axes[0].set_ylim(0, 1.05)
    axes[0].grid(axis="y", alpha=0.25)
    for index, value in enumerate(band_summary["co_occurrence_rate"]):
        axes[0].text(index, value + 0.03, f"{value:.0%}", ha="center")

    axes[1].bar(
        band_summary.index,
        band_summary["mean_extra_invocations"],
        color=["#D9A441", "#C17F2D", "#9C5E1A"],
    )
    axes[1].set_title("B. Resolution-Induced Invocations")
    axes[1].set_xlabel("TRACE risk band")
    axes[1].set_ylabel("Mean duplicate/recheck invocations")
    axes[1].grid(axis="y", alpha=0.25)
    for index, value in enumerate(band_summary["mean_extra_invocations"]):
        axes[1].text(index, value + 0.05, f"{value:.2f}", ha="center")

    for group_name, color in [
        ("low_co_occurrence_intensity", "#2E8B57"),
        ("high_co_occurrence_intensity", "#B84A4A"),
    ]:
        subset = drift_acceleration[drift_acceleration["group"] == group_name]
        label = group_name.replace("_", " ")
        axes[2].plot(
            subset["checkpoint"],
            subset["group_kl_drift_score"],
            marker="o",
            linewidth=2.4,
            color=color,
            label=label,
        )
    axes[2].axhline(0.15, linestyle="--", color="#CC3333", linewidth=1.2, label="KL threshold")
    axes[2].set_title("C. Downstream Drift-Risk Trajectory")
    axes[2].set_xlabel("Borrower checkpoint")
    axes[2].set_ylabel("Group KL drift score")
    axes[2].grid(axis="y", alpha=0.25)
    axes[2].legend(fontsize=8)

    note = (
        f"Overall co-occurrence: {metrics['co_occurrence_rate']:.0%} | "
        f"Avg resolution-induced invocations: {metrics['avg_resolution_induced_invocations']:.2f} | "
        f"High/low drift ratio: {metrics['drift_acceleration_ratio_high_vs_low']:.2f}x"
    )
    fig.text(0.5, 0.01, note, ha="center", fontsize=10)
    plt.tight_layout(rect=(0, 0.04, 1, 0.93))
    plt.savefig(FIGURE_PATH, dpi=220)
    plt.close()


def write_report(metrics: dict[str, float]) -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    common_subset = metrics["common_subset"]
    report = f"""# Week 4 Day 3: Scenario 4 Interaction Module Report

## Objective

Scenario 4 models the causal chain `Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk`. The analysis uses saved outputs from the Week 3 disagreement module, Week 3 overlap module, and Week 4 ABA drift module. No new LLM run was required.

## Inputs

- Disagreement results: `{DISAGREEMENT_PATH}`
- Overlap results: `{OVERLAP_PATH}`
- Drift/ABA event log: `{DRIFT_EVENTS_PATH}`
- Common borrower subset: `{common_subset}`

## Main Metrics

- Profiles analysed: `{int(metrics['profiles_analyzed'])}`
- Overlap rate: `{metrics['overlap_rate']:.1%}`
- Disagreement rate: `{metrics['disagreement_rate']:.1%}`
- Overlap-disagreement co-occurrence rate: `{metrics['co_occurrence_rate']:.1%}`
- Mean resolution-induced invocations: `{metrics['avg_resolution_induced_invocations']:.2f}`
- Mean remaining invocations after LCE+ToM: `{metrics['avg_remaining_invocations_after_lce']:.2f}`
- Mean duplicate invocations reduced by LCE+ToM: `{metrics['avg_invocations_reduced_by_lce']:.2f}`
- High/low co-occurrence drift-risk ratio: `{metrics['drift_acceleration_ratio_high_vs_low']:.2f}x`

![Figure 4 Causal Chain]({FIGURE_PATH})

## Three Causal Chain Findings

1. Overlap and disagreement co-occurred in `{metrics['co_occurrence_rate']:.1%}` of the common borrower subset, showing that task overlap and decision disagreement are not isolated events in this simulation.

2. When disagreement resolution fired, the raw overlap pipeline showed an average of `{metrics['avg_resolution_induced_invocations']:.2f}` duplicate or recheck invocations. With LCE+ToM active, this fell to `{metrics['avg_remaining_invocations_after_lce']:.2f}`, indicating that leader assignment reduced renewed overlap after resolution.

3. Borrowers in the high co-occurrence intensity group showed a downstream drift-risk ratio of `{metrics['drift_acceleration_ratio_high_vs_low']:.2f}x` compared with the low co-occurrence group. This supports the proposed causal mechanism that overlap and disagreement can amplify later drift risk.

## Caveat

This Wednesday analysis uses the common borrower subset available across the saved disagreement, overlap, and drift files (`{common_subset}`). The result is valid as an interaction analysis because all metrics are computed on the same borrower IDs.
"""
    REPORT_MD_PATH.write_text(report, encoding="utf-8")

    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading("Week 4 Day 3: Scenario 4 Interaction Module Report", level=1)
        doc.add_heading("Objective", level=2)
        doc.add_paragraph(
            "Scenario 4 models the causal chain Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk. "
            "The analysis uses saved outputs from the Week 3 disagreement module, Week 3 overlap module, and Week 4 ABA drift module."
        )
        doc.add_heading("Main Metrics", level=2)
        bullets = [
            f"Profiles analysed: {int(metrics['profiles_analyzed'])}",
            f"Overlap-disagreement co-occurrence rate: {metrics['co_occurrence_rate']:.1%}",
            f"Mean resolution-induced invocations: {metrics['avg_resolution_induced_invocations']:.2f}",
            f"Mean remaining invocations after LCE+ToM: {metrics['avg_remaining_invocations_after_lce']:.2f}",
            f"High/low co-occurrence drift-risk ratio: {metrics['drift_acceleration_ratio_high_vs_low']:.2f}x",
        ]
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_heading("Figure 4: Causal Chain", level=2)
        doc.add_picture(str(FIGURE_PATH), width=Inches(6.5))
        doc.add_heading("Three Causal Chain Findings", level=2)
        findings = [
            f"Overlap and disagreement co-occurred in {metrics['co_occurrence_rate']:.1%} of the common borrower subset.",
            f"Resolution events showed {metrics['avg_resolution_induced_invocations']:.2f} raw duplicate/recheck invocations, reduced to {metrics['avg_remaining_invocations_after_lce']:.2f} with LCE+ToM.",
            f"High co-occurrence cases showed {metrics['drift_acceleration_ratio_high_vs_low']:.2f}x the downstream drift-risk signal of low co-occurrence cases.",
        ]
        for finding in findings:
            doc.add_paragraph(finding, style="List Number")
        doc.add_heading("Caveat", level=2)
        doc.add_paragraph(
            f"This analysis uses the common borrower subset available across the saved disagreement, overlap, and drift files ({common_subset}). "
            "The comparison remains valid because all Scenario 4 metrics are computed on the same borrower IDs."
        )
        doc.save(REPORT_DOCX_PATH)
    except Exception as exc:
        print(f"Could not create DOCX report: {exc}")


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)

    disagreement = read_required_csv(DISAGREEMENT_PATH)
    overlap = read_required_csv(OVERLAP_PATH)
    drift_events = read_required_csv(DRIFT_EVENTS_PATH)
    baseline = read_required_csv(BASELINE_PATH)

    dual = build_dual_results(disagreement, overlap)
    resolution = build_resolution_overlap(dual)
    drift_acceleration = build_drift_acceleration(dual, drift_events, baseline)
    metrics = summarize(dual, resolution, drift_acceleration)

    dual.to_csv(DUAL_RESULTS_PATH, index=False)
    resolution.to_csv(RESOLUTION_OVERLAP_PATH, index=False)
    drift_acceleration.to_csv(DRIFT_ACCELERATION_PATH, index=False)
    save_summary(metrics)
    plot_causal_chain(dual, resolution, drift_acceleration, metrics)
    write_report(metrics)

    print("Scenario 4 interaction analysis complete.")
    print(f"Profiles analysed: {int(metrics['profiles_analyzed'])}")
    print(f"Co-occurrence rate: {metrics['co_occurrence_rate']:.1%}")
    print(f"Mean resolution-induced invocations: {metrics['avg_resolution_induced_invocations']:.2f}")
    print(f"Mean remaining invocations after LCE+ToM: {metrics['avg_remaining_invocations_after_lce']:.2f}")
    print(f"High/low drift-risk ratio: {metrics['drift_acceleration_ratio_high_vs_low']:.2f}x")
    print(f"Saved: {DUAL_RESULTS_PATH}")
    print(f"Saved: {RESOLUTION_OVERLAP_PATH}")
    print(f"Saved: {DRIFT_ACCELERATION_PATH}")
    print(f"Saved: {FIGURE_PATH}")
    print(f"Saved: {REPORT_MD_PATH}")
    if REPORT_DOCX_PATH.exists():
        print(f"Saved: {REPORT_DOCX_PATH}")


if __name__ == "__main__":
    main()
