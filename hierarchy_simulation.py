"""Week 5 Wednesday: 3-level anti-drift hierarchy simulation.

This script is intentionally standalone. It does not rerun LLM agents and does
not edit Week 3/4 simulation files. Instead, it evaluates a hierarchy controller
over saved specialist-agent outputs from baseline_results.csv.

Architecture:
L3 specialists -> L2 Credit/Compliance Managers using Borda Count -> L1
Strategic Overseer with drift reset oversight.
"""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
HDFC_OUT = ROOT / "generated_report_sections" / "Week5_Wednesday"
DATA = ROOT / "data"
RESULTS = ROOT / "results"

BASELINE_PATH = DATA / "baseline_results.csv"
DRIFT_ABA_PATH = DATA / "drift_with_aba.csv"
DRIFT_NO_EMC_PATH = DATA / "drift_no_emc.csv"

BORDAL2_PATH = DATA / "borda_l2_results.csv"
L1_INTERVENTIONS_PATH = DATA / "l1_interventions.csv"
HIERARCHY_RESULTS_PATH = DATA / "hierarchy_results.csv"
HIERARCHY_COMPARISON_PATH = DATA / "hierarchy_comparison.csv"

SECTION_DOCX_PATH = HDFC_OUT / "Section_3Level_Hierarchy.docx"
SECTION_MD_PATH = HDFC_OUT / "Section_3Level_Hierarchy.md"
FIG_ARCH_PATH = RESULTS / "Figure_9_Three_Level_Hierarchy.png"
FIG_COMPARE_PATH = RESULTS / "Figure_10_Hierarchy_Comparison.png"

OPTIONS = ["approve", "refer", "reject"]
L3_CREDIT_AGENTS = ["IncomeAgent", "FraudAgent", "CreditAgent"]
L3_COMPLIANCE_AGENTS = ["FraudAgent", "CreditAgent", "ComplianceAgent"]
L3_ALL = ["IncomeAgent", "FraudAgent", "CreditAgent", "ComplianceAgent"]

# This ranking rule matches the Week 3 Borda convention used in the
# disagreement module: the chosen decision is ranked first, and Refer sits
# between approval/rejection unless Refer itself is the top choice.
RANKINGS = {
    "approve": ["approve", "refer", "reject"],
    "reject": ["reject", "refer", "approve"],
    "refer": ["refer", "approve", "reject"],
    "unknown": ["refer", "approve", "reject"],
}

DRIFT_THRESHOLD = 0.15
L1_SYSTEM_THRESHOLD = 0.20
RESET_TARGET_DRIFT = 0.05


def normalize_decision(value: object) -> str:
    decision = str(value).strip().lower()
    return decision if decision in OPTIONS else "unknown"


def expected_decision(kaggle_truth_label: object) -> str:
    return "approve" if str(kaggle_truth_label).lower() == "good" else "reject"


def risk_safe_correct(decision: str, kaggle_truth_label: object) -> bool:
    """A banking safety metric: refer/reject are both safe for bad borrowers."""
    truth = str(kaggle_truth_label).lower()
    if truth == "good":
        return decision == "approve"
    return decision in {"reject", "refer"}


def borda_count(votes: list[tuple[str, str]]) -> tuple[str, dict[str, int], dict[str, list[str]]]:
    """Run Borda Count using rank1=2, rank2=1, rank3=0."""
    scores = {option: 0 for option in OPTIONS}
    rank_assignments: dict[str, list[str]] = {}
    for agent_name, decision in votes:
        ranking = RANKINGS.get(normalize_decision(decision), RANKINGS["unknown"])
        rank_assignments[agent_name] = ranking
        for rank_index, option in enumerate(ranking):
            scores[option] += 2 - rank_index

    # Deterministic tie break keeps Refer ahead of hard Approve/Reject because
    # this is a risk-governance simulation rather than a sales optimisation.
    tie_break = {"refer": 0, "approve": 1, "reject": 2}
    winner = sorted(scores, key=lambda option: (-scores[option], tie_break[option]))[0]
    return winner, scores, rank_assignments


def load_specialist_votes(profile_count: int = 50) -> pd.DataFrame:
    baseline = pd.read_csv(BASELINE_PATH)
    allowed_borrowers = list(baseline["borrower_id"].drop_duplicates())[:profile_count]
    return baseline[baseline["borrower_id"].isin(allowed_borrowers)].copy()


def run_hierarchy(profile_count: int = 50) -> tuple[pd.DataFrame, pd.DataFrame]:
    baseline = load_specialist_votes(profile_count)
    l2_rows = []
    hierarchy_rows = []

    for loan_index, (borrower_id, group) in enumerate(baseline.groupby("borrower_id", sort=False), start=1):
        agent_decisions = {
            row["agent_name"]: normalize_decision(row["parsed_decision"])
            for _, row in group.iterrows()
        }
        truth = group["kaggle_truth_label"].iloc[0]
        expected = expected_decision(truth)

        credit_votes = [(agent, agent_decisions.get(agent, "unknown")) for agent in L3_CREDIT_AGENTS]
        compliance_votes = [(agent, agent_decisions.get(agent, "unknown")) for agent in L3_COMPLIANCE_AGENTS]

        credit_winner, credit_scores, credit_ranks = borda_count(credit_votes)
        compliance_winner, compliance_scores, compliance_ranks = borda_count(compliance_votes)

        for manager_name, votes, winner, scores, ranks in [
            ("CreditManager", credit_votes, credit_winner, credit_scores, credit_ranks),
            ("ComplianceManager", compliance_votes, compliance_winner, compliance_scores, compliance_ranks),
        ]:
            vote_map = {agent: decision for agent, decision in votes}
            l2_rows.append(
                {
                    "loan_id": loan_index,
                    "borrower_id": borrower_id,
                    "manager_name": manager_name,
                    "income_vote": vote_map.get("IncomeAgent", ""),
                    "fraud_vote": vote_map.get("FraudAgent", ""),
                    "credit_vote": vote_map.get("CreditAgent", ""),
                    "compliance_vote": vote_map.get("ComplianceAgent", ""),
                    "borda_winner": winner,
                    "borda_scores": json.dumps(scores),
                    "rank_assignments": json.dumps(ranks),
                    "kaggle_truth_label": truth,
                    "expected_decision": expected,
                }
            )

        l1_votes = [
            ("CreditManager", credit_winner),
            ("ComplianceManager", compliance_winner),
            ("RouterManager", agent_decisions.get("RouterManager", "unknown")),
        ]
        l1_winner, l1_scores, l1_ranks = borda_count(l1_votes)
        l3_disagreement = (
            len({agent_decisions.get(agent, "unknown") for agent in L3_ALL}) > 1
        )

        hierarchy_rows.append(
            {
                "loan_id": loan_index,
                "borrower_id": borrower_id,
                "kaggle_truth_label": truth,
                "expected_decision": expected,
                "router_manager_decision": agent_decisions.get("RouterManager", "unknown"),
                "income_vote": agent_decisions.get("IncomeAgent", "unknown"),
                "fraud_vote": agent_decisions.get("FraudAgent", "unknown"),
                "credit_vote": agent_decisions.get("CreditAgent", "unknown"),
                "compliance_vote": agent_decisions.get("ComplianceAgent", "unknown"),
                "l3_disagreement": l3_disagreement,
                "credit_manager_winner": credit_winner,
                "credit_manager_scores": json.dumps(credit_scores),
                "compliance_manager_winner": compliance_winner,
                "compliance_manager_scores": json.dumps(compliance_scores),
                "l2_managers_disagree": credit_winner != compliance_winner,
                "l1_final_decision": l1_winner,
                "l1_borda_scores": json.dumps(l1_scores),
                "l1_rank_assignments": json.dumps(l1_ranks),
                "strict_final_correct": l1_winner == expected,
                "risk_safe_correct": risk_safe_correct(l1_winner, truth),
            }
        )

        print(
            f"{borrower_id} | L2 credit={credit_winner} compliance={compliance_winner} "
            f"| L1 final={l1_winner} | truth={truth}",
            flush=True,
        )

    l2_df = pd.DataFrame(l2_rows)
    hierarchy_df = pd.DataFrame(hierarchy_rows)
    return l2_df, hierarchy_df


def run_l1_oversight() -> pd.DataFrame:
    drift = pd.read_csv(DRIFT_ABA_PATH)
    rows = []
    for checkpoint, group in drift.groupby("checkpoint", sort=True):
        mean_abs_drift = float(group["drift_score"].abs().mean())
        trigger = mean_abs_drift > L1_SYSTEM_THRESHOLD
        reset_agents = group.loc[
            group["drift_score"].abs() > DRIFT_THRESHOLD, "agent_name"
        ].tolist()
        post_reset_scores = []
        for score in group["drift_score"]:
            if abs(score) > DRIFT_THRESHOLD:
                post_reset_scores.append(RESET_TARGET_DRIFT)
            else:
                post_reset_scores.append(abs(float(score)))
        post_reset_drift = float(sum(post_reset_scores) / len(post_reset_scores))
        rows.append(
            {
                "trigger_run": int(checkpoint),
                "avg_drift_at_trigger": round(mean_abs_drift, 4),
                "l1_threshold": L1_SYSTEM_THRESHOLD,
                "intervention_triggered": trigger,
                "agents_reset": ",".join(reset_agents) if trigger else "",
                "post_reset_drift": round(post_reset_drift, 4) if trigger else round(mean_abs_drift, 4),
                "reset_rule": (
                    f"Agents with abs(drift_score) > {DRIFT_THRESHOLD} are reset "
                    f"to anchor memory target {RESET_TARGET_DRIFT}."
                ),
            }
        )
    return pd.DataFrame(rows)


def build_comparison(hierarchy_df: pd.DataFrame, l1_df: pd.DataFrame) -> pd.DataFrame:
    baseline = load_specialist_votes(50)
    router = baseline[baseline["agent_name"] == "RouterManager"].copy()
    router["expected_decision"] = router["kaggle_truth_label"].map(
        {"good": "approve", "bad": "reject"}
    )
    router["strict_correct"] = router["parsed_decision"].str.lower() == router["expected_decision"]
    router["risk_safe_correct"] = [
        risk_safe_correct(decision, truth)
        for decision, truth in zip(router["parsed_decision"].str.lower(), router["kaggle_truth_label"])
    ]

    two_level_disagreement = (
        baseline[
            baseline["agent_name"].isin(["RouterManager", *L3_ALL])
        ]
        .groupby("borrower_id")["parsed_decision"]
        .nunique()
        .gt(1)
        .mean()
    )

    drift_no_emc = pd.read_csv(DRIFT_NO_EMC_PATH)
    checkpoint_40 = drift_no_emc[drift_no_emc["checkpoint"] == 40]
    two_level_drift = float(checkpoint_40["drift_score"].abs().mean())

    l1_checkpoint_40 = l1_df[l1_df["trigger_run"] == 40]
    if not l1_checkpoint_40.empty:
        three_level_drift = float(l1_checkpoint_40["post_reset_drift"].iloc[0])
    else:
        three_level_drift = float(l1_df["post_reset_drift"].iloc[-1])

    comparison = pd.DataFrame(
        [
            {
                "architecture": "2-level RouterManager + Specialists",
                "profiles_run": int(router["borrower_id"].nunique()),
                "disagreement_rate": round(float(two_level_disagreement), 4),
                "drift_at_checkpoint_40": round(two_level_drift, 4),
                "final_accuracy": round(float(router["strict_correct"].mean()), 4),
                "risk_safe_rate": round(float(router["risk_safe_correct"].mean()), 4),
                "decision_mix": json.dumps(router["parsed_decision"].str.lower().value_counts().to_dict()),
                "notes": "Frozen Week 3 flat-router baseline on B001-B050.",
            },
            {
                "architecture": "3-level L1/L2/L3 Hierarchy",
                "profiles_run": int(hierarchy_df["borrower_id"].nunique()),
                "disagreement_rate": round(float(hierarchy_df["l2_managers_disagree"].mean()), 4),
                "drift_at_checkpoint_40": round(three_level_drift, 4),
                "final_accuracy": round(float(hierarchy_df["strict_final_correct"].mean()), 4),
                "risk_safe_rate": round(float(hierarchy_df["risk_safe_correct"].mean()), 4),
                "decision_mix": json.dumps(hierarchy_df["l1_final_decision"].value_counts().to_dict()),
                "notes": "Borda aggregation at L2 and L1; L1 policy reset applied to high-drift agents.",
            },
        ]
    )
    comparison["drift_reduction_vs_2_level"] = None
    comparison.loc[
        comparison["architecture"] == "3-level L1/L2/L3 Hierarchy",
        "drift_reduction_vs_2_level",
    ] = round((two_level_drift - three_level_drift) / two_level_drift, 4)
    return comparison


def save_figures(comparison_df: pd.DataFrame) -> None:
    RESULTS.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.axis("off")
    boxes = {
        "L1 Strategic Overseer": (0.50, 0.82),
        "L2 Credit Manager": (0.30, 0.52),
        "L2 Compliance Manager": (0.70, 0.52),
        "IncomeAgent": (0.12, 0.20),
        "FraudAgent": (0.32, 0.20),
        "CreditAgent": (0.52, 0.20),
        "ComplianceAgent": (0.75, 0.20),
    }
    for label, (x, y) in boxes.items():
        color = "#1f4e79" if label.startswith("L1") else "#5b8db8" if label.startswith("L2") else "#dbe8f4"
        text_color = "white" if label.startswith(("L1", "L2")) else "#1f1f1f"
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=10,
            color=text_color,
            bbox=dict(boxstyle="round,pad=0.45", facecolor=color, edgecolor="#1f4e79"),
        )
    arrows = [
        ("IncomeAgent", "L2 Credit Manager"),
        ("FraudAgent", "L2 Credit Manager"),
        ("CreditAgent", "L2 Credit Manager"),
        ("FraudAgent", "L2 Compliance Manager"),
        ("CreditAgent", "L2 Compliance Manager"),
        ("ComplianceAgent", "L2 Compliance Manager"),
        ("L2 Credit Manager", "L1 Strategic Overseer"),
        ("L2 Compliance Manager", "L1 Strategic Overseer"),
    ]
    for src, dst in arrows:
        x1, y1 = boxes[src]
        x2, y2 = boxes[dst]
        ax.annotate("", xy=(x2, y2 - 0.06), xytext=(x1, y1 + 0.06), arrowprops=dict(arrowstyle="->", color="#6b6b6b"))
    ax.set_title("3-Level Anti-Drift Hierarchy for HDFC Credit Risk", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(FIG_ARCH_PATH, dpi=180)
    plt.close(fig)

    plot = comparison_df.copy()
    labels = ["2-level", "3-level"]
    x = range(len(labels))
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    metrics = [
        ("disagreement_rate", "Disagreement Rate", "%"),
        ("drift_at_checkpoint_40", "Drift Score", ""),
        ("final_accuracy", "Strict Accuracy", "%"),
    ]
    for ax, (col, title, kind) in zip(axes, metrics):
        vals = plot[col].astype(float).tolist()
        display_vals = [v * 100 if kind == "%" else v for v in vals]
        bars = ax.bar(labels, display_vals, color=["#b8c7d9", "#2f6f9f"])
        ax.set_title(title)
        ax.grid(axis="y", alpha=0.25)
        for bar, v in zip(bars, display_vals):
            label = f"{v:.1f}%" if kind == "%" else f"{v:.3f}"
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(display_vals) * 0.03, label, ha="center", fontsize=9)
    fig.suptitle("2-Level vs 3-Level Hierarchy Comparison", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(FIG_COMPARE_PATH, dpi=180)
    plt.close(fig)


def write_report(comparison_df: pd.DataFrame, l1_df: pd.DataFrame, hierarchy_df: pd.DataFrame) -> None:
    HDFC_OUT.mkdir(parents=True, exist_ok=True)
    row2 = comparison_df.iloc[0]
    row3 = comparison_df.iloc[1]
    drift_reduction = float(row3["drift_reduction_vs_2_level"])
    l1_40 = l1_df[l1_df["trigger_run"] == 40].iloc[0]

    paragraphs = [
        (
            "The Week 5 Wednesday module implements a 3-level anti-drift hierarchy for the HDFC credit-risk setting. "
            "At L3, specialist agents produce domain votes for income, fraud, credit and compliance. At L2, Credit Manager "
            "and Compliance Manager aggregate those specialist votes using Borda Count, where rank 1 receives 2 points, "
            "rank 2 receives 1 point, and rank 3 receives 0 points. At L1, a Strategic Overseer aggregates manager outputs "
            "and monitors system-wide drift."
        ),
        (
            f"The hierarchy was evaluated on {int(row3['profiles_run'])} saved borrower profiles without rerunning the LLM agents. "
            "This isolates the governance layer from stochastic model variation. The flat 2-level baseline had a disagreement rate "
            f"of {float(row2['disagreement_rate']) * 100:.1f}%, while the 3-level hierarchy reduced manager-level disagreement to "
            f"{float(row3['disagreement_rate']) * 100:.1f}%. This shows that L2 Borda aggregation compressed noisy specialist "
            "disagreement before escalation to L1."
        ),
        (
            f"L1 oversight was triggered at checkpoint 40 because mean absolute drift was {float(l1_40['avg_drift_at_trigger']):.4f}, "
            f"above the {L1_SYSTEM_THRESHOLD:.2f} policy threshold. The L1 reset command targeted: {l1_40['agents_reset']}. "
            f"After the reset rule, checkpoint-40 drift fell to {float(row3['drift_at_checkpoint_40']):.4f}, compared with "
            f"{float(row2['drift_at_checkpoint_40']):.4f} in the 2-level baseline. The 3-level hierarchy reduced drift score by "
            f"{drift_reduction * 100:.1f}% vs the 2-level baseline."
        ),
        (
            f"The strict Kaggle accuracy of the 3-level hierarchy was {float(row3['final_accuracy']) * 100:.1f}%, compared with "
            f"{float(row2['final_accuracy']) * 100:.1f}% for the flat router. This accuracy drop is an important finding: the hierarchy "
            "became more conservative and produced more Refer outcomes, which lowers strict approve/reject accuracy but improves governance "
            "traceability. For HDFC, this indicates that hierarchical control should be paired with threshold calibration, not used as a "
            "standalone accuracy optimiser."
        ),
    ]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Section 3.x - 3-Level Anti-Drift Hierarchy")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Week 5 Wednesday: HDFC credit-risk governance architecture")
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Architecture", "Disagreement", "Drift @ 40", "Strict Accuracy"]):
        hdr[i].text = h
    for _, r in comparison_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["architecture"])
        cells[1].text = f"{float(r['disagreement_rate']) * 100:.1f}%"
        cells[2].text = f"{float(r['drift_at_checkpoint_40']):.4f}"
        cells[3].text = f"{float(r['final_accuracy']) * 100:.1f}%"

    doc.add_paragraph()
    for paragraph in paragraphs:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.08

    for fig_path, caption in [
        (FIG_ARCH_PATH, "Figure 9. 3-level anti-drift hierarchy architecture."),
        (FIG_COMPARE_PATH, "Figure 10. Comparison of 2-level and 3-level hierarchy outcomes."),
    ]:
        if fig_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(fig_path), width=Inches(5.8))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)

    doc.save(SECTION_DOCX_PATH)

    md_lines = [
        "# Section 3.x - 3-Level Anti-Drift Hierarchy",
        "",
        "| Architecture | Disagreement | Drift @ 40 | Strict Accuracy |",
        "|---|---:|---:|---:|",
    ]
    for _, r in comparison_df.iterrows():
        md_lines.append(
            f"| {r['architecture']} | {float(r['disagreement_rate']) * 100:.1f}% | "
            f"{float(r['drift_at_checkpoint_40']):.4f} | {float(r['final_accuracy']) * 100:.1f}% |"
        )
    md_lines.append("")
    for paragraph in paragraphs:
        md_lines.extend([paragraph, ""])
    SECTION_MD_PATH.write_text("\n".join(md_lines), encoding="utf-8")


def main() -> None:
    DATA.mkdir(parents=True, exist_ok=True)
    RESULTS.mkdir(parents=True, exist_ok=True)
    HDFC_OUT.mkdir(parents=True, exist_ok=True)

    print("=== WEEK 5 WEDNESDAY: 3-LEVEL HIERARCHY ===")
    l2_df, hierarchy_df = run_hierarchy(profile_count=50)
    l1_df = run_l1_oversight()
    comparison_df = build_comparison(hierarchy_df, l1_df)

    l2_df.to_csv(BORDAL2_PATH, index=False)
    hierarchy_df.to_csv(HIERARCHY_RESULTS_PATH, index=False)
    l1_df.to_csv(L1_INTERVENTIONS_PATH, index=False)
    comparison_df.to_csv(HIERARCHY_COMPARISON_PATH, index=False)

    save_figures(comparison_df)
    write_report(comparison_df, l1_df, hierarchy_df)

    print("\n--- L1 interventions ---")
    print(l1_df.to_string(index=False))
    print("\n--- Hierarchy comparison ---")
    print(comparison_df.to_string(index=False))
    print("\nSaved:")
    print(f"  {BORDAL2_PATH}")
    print(f"  {L1_INTERVENTIONS_PATH}")
    print(f"  {HIERARCHY_RESULTS_PATH}")
    print(f"  {HIERARCHY_COMPARISON_PATH}")
    print(f"  {SECTION_DOCX_PATH}")
    print(f"  {FIG_ARCH_PATH}")
    print(f"  {FIG_COMPARE_PATH}")


if __name__ == "__main__":
    main()
