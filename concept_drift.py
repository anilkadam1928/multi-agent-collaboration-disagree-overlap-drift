from __future__ import annotations

"""Week 5 Monday - Concept Drift Correlation Layer.

This script is intentionally separate from the existing simulation files.
It does not call the LLM. It measures whether incoming borrower batches look
statistically different from the frozen first-50 German Credit baseline.

Method:
    1. Use B001-B050 as the frozen baseline distribution.
    2. Split B001-B200 into four 50-profile batches.
    3. Convert each batch into a distribution vector:
       - numeric fields become quantile-bin proportions
       - categorical fields become category proportions
    4. Compute Pearson correlation r between each batch vector and baseline.
    5. If r < 0.70, flag concept drift and recommend recalibration.

Outputs:
    data/concept_drift_log.csv
    data/concept_drift_feature_summary.csv
    results/Figure_6_Concept_Drift_Correlation.png
    generated_report_sections/Week5_Monday/Section_3_ConceptDrift_Addition.docx
"""

import csv
import math
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent
HDFC_OUTPUT_DIR = PROJECT_ROOT / "generated_report_sections" / "Week5_Monday"

DATASET_PATH = PROJECT_ROOT / "german_credit_data.csv"
LOG_PATH = PROJECT_ROOT / "data" / "concept_drift_log.csv"
FEATURE_SUMMARY_PATH = PROJECT_ROOT / "data" / "concept_drift_feature_summary.csv"
FIGURE_PATH = PROJECT_ROOT / "results" / "Figure_6_Concept_Drift_Correlation.png"
SECTION3_DOCX_PATH = HDFC_OUTPUT_DIR / "Section_3_ConceptDrift_Addition.docx"
SECTION3_MD_PATH = HDFC_OUTPUT_DIR / "Section_3_ConceptDrift_Addition.md"

BASELINE_SIZE = 50
EVAL_PROFILE_COUNT = 200
BATCH_SIZE = 50
CONCEPT_DRIFT_THRESHOLD = 0.70

NUMERIC_COLUMNS = ["Age", "Job", "Credit amount", "Duration"]
CATEGORICAL_COLUMNS = [
    "Sex",
    "Housing",
    "Saving accounts",
    "Checking account",
    "Purpose",
]


def clean_category(value: object) -> str:
    text = str(value).strip().lower()
    if text in {"", "nan", "na", "n/a", "none", "null"}:
        return "unknown"
    return text


def load_dataset() -> pd.DataFrame:
    if not DATASET_PATH.exists():
        raise FileNotFoundError(f"Missing dataset: {DATASET_PATH}")
    df = pd.read_csv(DATASET_PATH)
    df = df.drop(columns=["Unnamed: 0"], errors="ignore")
    for column in NUMERIC_COLUMNS:
        df[column] = pd.to_numeric(df[column], errors="coerce")
    for column in CATEGORICAL_COLUMNS:
        df[column] = df[column].map(clean_category)
    return df


def build_numeric_bins(baseline: pd.DataFrame) -> dict[str, list[float]]:
    bins: dict[str, list[float]] = {}
    for column in NUMERIC_COLUMNS:
        values = baseline[column].dropna()
        quantiles = values.quantile([0, 0.2, 0.4, 0.6, 0.8, 1.0]).tolist()
        unique_edges = sorted(set(float(value) for value in quantiles))
        if len(unique_edges) < 3:
            min_value = float(values.min())
            max_value = float(values.max())
            midpoint = (min_value + max_value) / 2
            unique_edges = [min_value, midpoint, max_value]
        unique_edges[0] = -math.inf
        unique_edges[-1] = math.inf
        bins[column] = unique_edges
    return bins


def distribution_vector(
    frame: pd.DataFrame,
    numeric_bins: dict[str, list[float]],
    category_levels: dict[str, list[str]],
) -> dict[str, float]:
    vector: dict[str, float] = {}
    total = max(len(frame), 1)

    for column in NUMERIC_COLUMNS:
        binned = pd.cut(
            frame[column],
            bins=numeric_bins[column],
            include_lowest=True,
            duplicates="drop",
        )
        proportions = binned.value_counts(normalize=True, sort=False)
        for index, interval in enumerate(proportions.index):
            vector[f"{column}_bin_{index + 1}"] = float(proportions.get(interval, 0.0))

    for column in CATEGORICAL_COLUMNS:
        counts = frame[column].map(clean_category).value_counts(normalize=True)
        for level in category_levels[column]:
            vector[f"{column}_{level}"] = float(counts.get(level, 0.0))

    # Include risk-label distribution separately because it is useful for analysis,
    # but in production this would be replaced by delayed observed outcomes.
    if "Risk" in frame.columns:
        risk_counts = frame["Risk"].map(clean_category).value_counts(normalize=True)
        vector["observed_risk_good"] = float(risk_counts.get("good", 0.0))
        vector["observed_risk_bad"] = float(risk_counts.get("bad", 0.0))

    # Guard against accidental empty batches.
    return {key: value if not math.isnan(value) else 0.0 for key, value in vector.items()}


def pearson_correlation(left: dict[str, float], right: dict[str, float]) -> float:
    keys = sorted(set(left) | set(right))
    x = [left.get(key, 0.0) for key in keys]
    y = [right.get(key, 0.0) for key in keys]
    mean_x = sum(x) / len(x)
    mean_y = sum(y) / len(y)
    numerator = sum((a - mean_x) * (b - mean_y) for a, b in zip(x, y))
    denom_x = math.sqrt(sum((a - mean_x) ** 2 for a in x))
    denom_y = math.sqrt(sum((b - mean_y) ** 2 for b in y))
    if denom_x == 0 or denom_y == 0:
        return 0.0
    return numerator / (denom_x * denom_y)


def make_stress_batch(batch: pd.DataFrame) -> pd.DataFrame:
    """Create a controlled stress case if natural batches do not cross threshold."""
    shifted = batch.copy()
    shifted["Credit amount"] = shifted["Credit amount"] * 2.5
    shifted["Duration"] = shifted["Duration"] * 1.8
    shifted["Age"] = (shifted["Age"] + 15).clip(upper=80)
    shifted["Checking account"] = "little"
    shifted["Saving accounts"] = "unknown"
    shifted["Purpose"] = "business"
    return shifted


def summarize_batch(batch: pd.DataFrame, batch_id: str, batch_type: str) -> dict[str, object]:
    return {
        "batch_id": batch_id,
        "batch_type": batch_type,
        "profiles_in_batch": len(batch),
        "age_mean": round(float(batch["Age"].mean()), 3),
        "credit_amount_mean": round(float(batch["Credit amount"].mean()), 3),
        "duration_mean": round(float(batch["Duration"].mean()), 3),
        "good_rate": round(float((batch["Risk"].map(clean_category) == "good").mean()), 3),
        "bad_rate": round(float((batch["Risk"].map(clean_category) == "bad").mean()), 3),
    }


def concept_drift_recommendation(drift_detected: bool, batch_type: str) -> str:
    if drift_detected and batch_type == "natural":
        return "Concept drift detected; recalibrate agents before relying on current decisions."
    if drift_detected:
        return "Stress-test drift detected; recalibration trigger behaves as expected."
    return "Stable; continue normal monitoring."


def run_concept_drift() -> tuple[pd.DataFrame, pd.DataFrame]:
    df = load_dataset().head(EVAL_PROFILE_COUNT).copy()
    baseline = df.head(BASELINE_SIZE).copy()

    numeric_bins = build_numeric_bins(baseline)
    category_levels = {
        column: sorted(df[column].map(clean_category).dropna().unique().tolist())
        for column in CATEGORICAL_COLUMNS
    }
    baseline_vector = distribution_vector(baseline, numeric_bins, category_levels)

    log_rows: list[dict[str, object]] = []
    summary_rows: list[dict[str, object]] = []

    for batch_index, start in enumerate(range(0, EVAL_PROFILE_COUNT, BATCH_SIZE), start=1):
        batch = df.iloc[start : start + BATCH_SIZE].copy()
        if batch.empty:
            continue
        batch_id = f"batch_{batch_index}_B{start + 1:03d}_B{start + len(batch):03d}"
        batch_vector = distribution_vector(batch, numeric_bins, category_levels)
        correlation = pearson_correlation(baseline_vector, batch_vector)
        drift_detected = correlation < CONCEPT_DRIFT_THRESHOLD
        log_rows.append(
            {
                "batch_id": batch_id,
                "batch_type": "natural",
                "profiles_in_batch": len(batch),
                "correlation_r": round(correlation, 4),
                "drift_detected": "Y" if drift_detected else "N",
                "threshold": CONCEPT_DRIFT_THRESHOLD,
                "recommendation": concept_drift_recommendation(drift_detected, "natural"),
            }
        )
        summary_rows.append(summarize_batch(batch, batch_id, "natural"))

    if not any(row["drift_detected"] == "Y" for row in log_rows):
        stress_source = df.iloc[150:200].copy()
        stress_batch = make_stress_batch(stress_source)
        stress_vector = distribution_vector(stress_batch, numeric_bins, category_levels)
        stress_r = pearson_correlation(baseline_vector, stress_vector)
        stress_detected = stress_r < CONCEPT_DRIFT_THRESHOLD
        stress_id = "stress_test_batch_4_shifted_credit_profile"
        log_rows.append(
            {
                "batch_id": stress_id,
                "batch_type": "synthetic_stress_test",
                "profiles_in_batch": len(stress_batch),
                "correlation_r": round(stress_r, 4),
                "drift_detected": "Y" if stress_detected else "N",
                "threshold": CONCEPT_DRIFT_THRESHOLD,
                "recommendation": concept_drift_recommendation(stress_detected, "synthetic_stress_test"),
            }
        )
        summary_rows.append(summarize_batch(stress_batch, stress_id, "synthetic_stress_test"))

    log_df = pd.DataFrame(log_rows)
    summary_df = pd.DataFrame(summary_rows)

    LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    FEATURE_SUMMARY_PATH.parent.mkdir(parents=True, exist_ok=True)
    log_df.to_csv(LOG_PATH, index=False)
    summary_df.to_csv(FEATURE_SUMMARY_PATH, index=False)
    return log_df, summary_df


def plot_concept_drift(log_df: pd.DataFrame) -> None:
    FIGURE_PATH.parent.mkdir(parents=True, exist_ok=True)
    colors = [
        "#B84A4A" if row["drift_detected"] == "Y" else "#2E8B57"
        for _, row in log_df.iterrows()
    ]
    labels = [str(batch_id).replace("_", "\n", 2) for batch_id in log_df["batch_id"]]

    plt.figure(figsize=(10, 5.2))
    plt.bar(labels, log_df["correlation_r"], color=colors)
    plt.axhline(CONCEPT_DRIFT_THRESHOLD, color="#CC3333", linestyle="--", linewidth=1.5)
    plt.ylim(0, 1.05)
    plt.ylabel("Pearson correlation r vs frozen B001-B050 baseline")
    plt.title("Figure 6 - Concept Drift Correlation Layer")
    plt.grid(axis="y", alpha=0.25)
    for index, value in enumerate(log_df["correlation_r"]):
        plt.text(index, float(value) + 0.025, f"{float(value):.2f}", ha="center")
    plt.tight_layout()
    plt.savefig(FIGURE_PATH, dpi=200)
    plt.close()


def build_section3_text(log_df: pd.DataFrame) -> str:
    natural = log_df[log_df["batch_type"] == "natural"]
    detected = natural[natural["drift_detected"] == "Y"]
    stress = log_df[log_df["batch_type"] == "synthetic_stress_test"]

    if not detected.empty:
        first = detected.iloc[0]
        detection_sentence = (
            f"{first['batch_id']} triggered drift at r={first['correlation_r']}, "
            "so the system recommended agent recalibration."
        )
    elif not stress.empty:
        first = stress.iloc[0]
        detection_sentence = (
            "No natural 50-profile batch crossed the drift threshold in the first 200 profiles. "
            f"A controlled stress test triggered the detector at r={first['correlation_r']}, "
            "confirming that the alert logic works when the incoming applicant distribution shifts."
        )
    else:
        min_row = natural.sort_values("correlation_r").iloc[0]
        detection_sentence = (
            f"No batch triggered drift; the lowest natural correlation was r={min_row['correlation_r']} "
            f"for {min_row['batch_id']}."
        )

    return (
        "A concept drift correlation layer was added to distinguish data-distribution shift from "
        "behavioural agent drift. The first 50 German Credit profiles (B001-B050) were treated as "
        "the frozen baseline distribution. Every new 50-profile batch was converted into a feature "
        "distribution vector using numeric-bin proportions for age, job, credit amount, and duration, "
        "plus category proportions for borrower attributes. Pearson correlation against the baseline "
        "was then computed, with r < 0.70 used as the concept drift threshold. "
        f"{detection_sentence} This layer prevents the system from wrongly treating all performance "
        "change as agent drift when the incoming borrower population itself has changed."
    )


def write_section3_doc(text: str) -> None:
    HDFC_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SECTION3_MD_PATH.write_text("# Section 3 Addition: Concept Drift Correlation Layer\n\n" + text + "\n", encoding="utf-8")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Section 3 Addition: Concept Drift Correlation Layer")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Week 5 Monday | Scenario 5 remaining component")
    sub.italic = True
    sub.font.size = Pt(9)
    sub.font.color.rgb = RGBColor(90, 90, 90)

    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.space_after = Pt(8)

    doc.save(SECTION3_DOCX_PATH)


def main() -> None:
    log_df, summary_df = run_concept_drift()
    plot_concept_drift(log_df)
    section_text = build_section3_text(log_df)
    write_section3_doc(section_text)

    print("Concept drift layer complete.")
    for _, row in log_df.iterrows():
        status = "CONCEPT DRIFT DETECTED" if row["drift_detected"] == "Y" else "stable"
        print(f"{row['batch_id']}: r={row['correlation_r']} ({status})")
    print(f"Saved log: {LOG_PATH}")
    print(f"Saved feature summary: {FEATURE_SUMMARY_PATH}")
    print(f"Saved figure: {FIGURE_PATH}")
    print(f"Saved Section 3 addition: {SECTION3_DOCX_PATH}")
    print(f"Saved Section 3 markdown: {SECTION3_MD_PATH}")


if __name__ == "__main__":
    main()
