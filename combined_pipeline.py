from __future__ import annotations

"""Week 4 Thursday - Scenario 5 integration layer.

This script does not run the local LLM. It integrates the saved outputs from:
- Baseline
- Disagreement module
- Overlap module
- Drift/ABA module
- Interaction module

It creates the Thursday deliverables:
- data/combined_results.csv
- analysis/master_comparison_table_week4.csv
- results/Figure_5_Master_Comparison.png
- data/rl_rewards.csv
- analysis/Week4_Day4_Combined_Report.docx
"""

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd


PROJECT_ROOT = Path(__file__).resolve().parent
DATA_DIR = PROJECT_ROOT / "data"
RESULTS_DIR = PROJECT_ROOT / "results"
ANALYSIS_DIR = PROJECT_ROOT / "analysis"

BASELINE_PATH = DATA_DIR / "baseline_results.csv"
DISAGREEMENT_PATH = RESULTS_DIR / "disagreement_results.csv"
OVERLAP_PATH = DATA_DIR / "overlap_results.csv"
DRIFT_WITH_EMC_PATH = DATA_DIR / "drift_with_emc.csv"
DRIFT_WITH_ABA_PATH = DATA_DIR / "drift_with_aba.csv"
DRIFT_ABA_EVENTS_PATH = DATA_DIR / "drift_with_aba_events.csv"
DUAL_RESULTS_PATH = DATA_DIR / "dual_results.csv"

COMBINED_RESULTS_PATH = DATA_DIR / "combined_results.csv"
MASTER_TABLE_CSV_PATH = ANALYSIS_DIR / "master_comparison_table_week4.csv"
MASTER_TABLE_MD_PATH = ANALYSIS_DIR / "master_comparison_table_week4.md"
RL_REWARDS_PATH = DATA_DIR / "rl_rewards.csv"
FIGURE_PATH = RESULTS_DIR / "Figure_5_Master_Comparison.png"
REPORT_MD_PATH = ANALYSIS_DIR / "Week4_Day4_Combined_Report.md"
REPORT_DOCX_PATH = ANALYSIS_DIR / "Week4_Day4_Combined_Report.docx"

VALID_FINAL_DECISIONS = {"approve", "reject"}


def read_required_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Missing required input file: {path}")
    return pd.read_csv(path)


def as_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"true", "1", "yes", "y"}


def borrower_number(series: pd.Series) -> pd.Series:
    return series.astype(str).str.extract(r"(\d+)").astype(int)[0]


def borrower_subset_label(borrower_ids: pd.Series | list[str]) -> str:
    series = pd.Series(borrower_ids, dtype=str)
    ordered = (
        series.drop_duplicates()
        .sort_values(key=lambda values: borrower_number(values))
        .tolist()
    )
    if not ordered:
        return "no common borrowers"
    if len(ordered) == 1:
        return ordered[0]
    return f"{ordered[0]}-{ordered[-1]}"


def expected_decision(kaggle_label: str) -> str:
    return "approve" if str(kaggle_label).lower() == "good" else "reject"


def decision_correct(decision: str, kaggle_label: str) -> bool:
    return str(decision).lower() == expected_decision(kaggle_label)


def reward_for_decision(decision: str, kaggle_label: str) -> int:
    decision = str(decision).lower()
    if decision_correct(decision, kaggle_label):
        return 1
    if decision in VALID_FINAL_DECISIONS:
        return -1
    return 0


def clean_decision(value) -> str:
    text = str(value).strip().lower()
    return text if text else "unknown"


def common_borrowers(*frames: pd.DataFrame) -> list[str]:
    borrower_sets = []
    for frame in frames:
        borrower_sets.append(set(frame["borrower_id"].astype(str)))
    common = set.intersection(*borrower_sets)
    return sorted(common, key=lambda bid: int("".join(ch for ch in bid if ch.isdigit()) or 0))


def baseline_table(baseline: pd.DataFrame, borrowers: list[str]) -> tuple[pd.DataFrame, dict[str, float]]:
    base_common = baseline[baseline["borrower_id"].isin(borrowers)].copy()
    router = base_common[base_common["agent_name"] == "RouterManager"].copy()
    router["baseline_router_decision"] = router["parsed_decision"].map(clean_decision)
    router["baseline_correct"] = router.apply(
        lambda row: decision_correct(row["baseline_router_decision"], row["kaggle_truth_label"]),
        axis=1,
    )

    disagreement_by_borrower = (
        base_common.groupby("borrower_id")["parsed_decision"]
        .apply(lambda values: len(set(clean_decision(value) for value in values)) > 1)
        .rename("baseline_agent_disagreement")
        .reset_index()
    )
    router = router.merge(disagreement_by_borrower, on="borrower_id", how="left")

    metrics = {
        "baseline_accuracy": float(router["baseline_correct"].mean()),
        "baseline_disagreement_rate": float(router["baseline_agent_disagreement"].mean()),
        "baseline_profiles": float(router["borrower_id"].nunique()),
    }
    return router[
        [
            "borrower_id",
            "kaggle_truth_label",
            "baseline_router_decision",
            "baseline_correct",
            "baseline_agent_disagreement",
        ]
    ], metrics


def disagreement_table(disagreement: pd.DataFrame, borrowers: list[str]) -> tuple[pd.DataFrame, dict[str, float]]:
    disagree_common = disagreement[disagreement["borrower_id"].isin(borrowers)].copy()
    disagree_common["disagreement_exists"] = disagree_common["disagreement_exists"].map(as_bool)
    disagree_common["sycophancy_flagged"] = disagree_common["sycophancy_flagged"].map(as_bool)
    disagree_common["position_changed"] = (
        disagree_common["precommit_decision"].map(clean_decision)
        != disagree_common["final_decision"].map(clean_decision)
    )

    rows = []
    for borrower_id, group in disagree_common.groupby("borrower_id", sort=True):
        first = group.iloc[0]
        borda_winner = clean_decision(first["borda_winner"])
        consens_winner = clean_decision(first["consensagent_winner"])
        rows.append(
            {
                "borrower_id": borrower_id,
                "disagreement_detected": bool(group["disagreement_exists"].max()),
                "borda_winner": borda_winner,
                "borda_correct": decision_correct(borda_winner, first["kaggle_truth_label"]),
                "consensagent_winner": consens_winner,
                "consensagent_correct": decision_correct(consens_winner, first["kaggle_truth_label"]),
                "position_change_count": int(group["position_changed"].sum()),
                "sycophancy_flag_count": int(group["sycophancy_flagged"].sum()),
            }
        )

    table = pd.DataFrame(rows)
    total_changes = int(disagree_common["position_changed"].sum())
    syco_flags = int(disagree_common["sycophancy_flagged"].sum())
    metrics = {
        "disagreement_rate": float(table["disagreement_detected"].mean()),
        "borda_accuracy": float(table["borda_correct"].mean()),
        "consensagent_accuracy": float(table["consensagent_correct"].mean()),
        "sycophancy_rate": float(syco_flags / total_changes) if total_changes else 0.0,
        "disagreement_profiles": float(table["borrower_id"].nunique()),
    }
    return table, metrics


def overlap_table(overlap: pd.DataFrame, borrowers: list[str]) -> tuple[pd.DataFrame, dict[str, float]]:
    overlap_common = overlap[overlap["borrower_id"].isin(borrowers)].copy()
    raw = overlap_common[overlap_common["mode"] == "without_lce"].copy()
    controlled = overlap_common[overlap_common["mode"] == "with_lce_tom"].copy()

    raw = raw.rename(
        columns={
            "redundancy_index": "raw_redundancy_index",
            "duplicate_invocations": "raw_duplicate_invocations",
            "tap_overlap_count": "raw_tap_overlap_count",
            "final_decision": "raw_overlap_final_decision",
            "final_correct": "raw_overlap_correct",
        }
    )
    controlled = controlled.rename(
        columns={
            "redundancy_index": "controlled_redundancy_index",
            "duplicate_invocations": "controlled_duplicate_invocations",
            "tap_overlap_count": "controlled_tap_overlap_count",
            "final_decision": "controlled_overlap_final_decision",
            "final_correct": "controlled_overlap_correct",
            "active_leader": "overlap_active_leader",
        }
    )

    table = raw[
        [
            "borrower_id",
            "trace_bid",
            "raw_redundancy_index",
            "raw_duplicate_invocations",
            "raw_tap_overlap_count",
            "raw_overlap_final_decision",
            "raw_overlap_correct",
        ]
    ].merge(
        controlled[
            [
                "borrower_id",
                "controlled_redundancy_index",
                "controlled_duplicate_invocations",
                "controlled_tap_overlap_count",
                "controlled_overlap_final_decision",
                "controlled_overlap_correct",
                "overlap_active_leader",
            ]
        ],
        on="borrower_id",
        how="inner",
    )
    table["raw_overlap_correct"] = table["raw_overlap_correct"].map(as_bool)
    table["controlled_overlap_correct"] = table["controlled_overlap_correct"].map(as_bool)
    table["redundancy_reduction"] = (
        table["raw_redundancy_index"] - table["controlled_redundancy_index"]
    ).clip(lower=0)
    table["overlap_detected"] = (
        (table["raw_redundancy_index"] > 0)
        | (table["raw_duplicate_invocations"] > 0)
        | (table["raw_tap_overlap_count"] > 0)
    )

    raw_mean = float(table["raw_redundancy_index"].mean())
    controlled_mean = float(table["controlled_redundancy_index"].mean())
    metrics = {
        "raw_redundancy_index": raw_mean,
        "controlled_redundancy_index": controlled_mean,
        "redundancy_reduction_percent": (1 - controlled_mean / raw_mean) * 100 if raw_mean else 0.0,
        "overlap_accuracy": float(table["controlled_overlap_correct"].mean()),
        "overlap_profiles": float(table["borrower_id"].nunique()),
    }
    return table, metrics


def drift_table(
    drift_events: pd.DataFrame,
    drift_with_emc: pd.DataFrame,
    drift_with_aba: pd.DataFrame,
    borrowers: list[str],
) -> tuple[pd.DataFrame, dict[str, float]]:
    events_common = drift_events[drift_events["borrower_id"].isin(borrowers)].copy()
    router = events_common[events_common["agent_name"] == "RouterManager"].copy()
    router["drift_router_decision"] = router["parsed_decision"].map(clean_decision)
    router["drift_router_correct"] = router.apply(
        lambda row: decision_correct(row["drift_router_decision"], row["kaggle_truth_label"]),
        axis=1,
    )
    router["aba_triggered"] = router["aba_triggered"].map(as_bool)

    latest_emc_checkpoint = int(drift_with_emc["checkpoint"].max())
    latest_aba_checkpoint = int(drift_with_aba["checkpoint"].max())
    emc_latest = drift_with_emc[drift_with_emc["checkpoint"] == latest_emc_checkpoint]
    aba_latest = drift_with_aba[drift_with_aba["checkpoint"] == latest_aba_checkpoint]

    with_emc_mean = float(emc_latest["drift_score"].mean())
    with_aba_mean = float(aba_latest["drift_score"].mean())
    metrics = {
        "with_emc_mean_drift": with_emc_mean,
        "with_aba_mean_drift": with_aba_mean,
        "aba_reduction_vs_emc_percent": (1 - with_aba_mean / with_emc_mean) * 100 if with_emc_mean else 0.0,
        "drift_router_accuracy_common": float(router["drift_router_correct"].mean()),
        "drift_profiles_available": float(drift_events["borrower_id"].nunique()),
        "drift_common_profiles": float(router["borrower_id"].nunique()),
    }
    return router[
        [
            "borrower_id",
            "drift_router_decision",
            "drift_router_correct",
            "pre_call_drift_score",
            "aba_triggered",
        ]
    ], metrics


def interaction_table(dual: pd.DataFrame, borrowers: list[str]) -> tuple[pd.DataFrame, dict[str, float]]:
    dual_common = dual[dual["borrower_id"].isin(borrowers)].copy()
    for column in ["co_occurrence", "borda_resolution_fired", "overlap_detected", "disagreement_detected"]:
        if column in dual_common.columns:
            dual_common[column] = dual_common[column].map(as_bool)

    metrics = {
        "co_occurrence_rate": float(dual_common["co_occurrence"].mean()),
        "avg_resolution_induced_invocations": float(dual_common["resolution_induced_invocations"].mean()),
        "avg_remaining_invocations_after_lce": float(dual_common["remaining_invocations_after_lce"].mean()),
        "interaction_profiles": float(dual_common["borrower_id"].nunique()),
    }
    return dual_common[
        [
            "borrower_id",
            "co_occurrence",
            "resolution_induced_invocations",
            "remaining_invocations_after_lce",
            "risk_band",
        ]
    ], metrics


def choose_combined_decision(row: pd.Series) -> tuple[str, str]:
    candidates = [
        ("Drift RouterManager", row.get("drift_router_decision")),
        ("Overlap LCE+ToM", row.get("controlled_overlap_final_decision")),
        ("Disagreement Borda", row.get("borda_winner")),
        ("Baseline RouterManager", row.get("baseline_router_decision")),
    ]
    for source, decision in candidates:
        decision = clean_decision(decision)
        if decision in VALID_FINAL_DECISIONS:
            return decision, source
    return "unknown", "No valid final decision"


def build_combined_results(
    baseline_part: pd.DataFrame,
    disagreement_part: pd.DataFrame,
    overlap_part: pd.DataFrame,
    drift_part: pd.DataFrame,
    interaction_part: pd.DataFrame,
) -> pd.DataFrame:
    combined = baseline_part.merge(disagreement_part, on="borrower_id", how="inner")
    combined = combined.merge(overlap_part, on="borrower_id", how="inner")
    combined = combined.merge(drift_part, on="borrower_id", how="inner")
    combined = combined.merge(interaction_part, on="borrower_id", how="inner")
    combined["run_id"] = combined["borrower_id"].map(lambda bid: f"combined_{bid}")

    decisions = combined.apply(choose_combined_decision, axis=1)
    combined["combined_final_decision"] = [item[0] for item in decisions]
    combined["combined_decision_source"] = [item[1] for item in decisions]
    combined["combined_final_correct"] = combined.apply(
        lambda row: decision_correct(row["combined_final_decision"], row["kaggle_truth_label"]),
        axis=1,
    )
    combined["combined_reward"] = combined.apply(
        lambda row: reward_for_decision(row["combined_final_decision"], row["kaggle_truth_label"]),
        axis=1,
    )
    return combined.sort_values("borrower_id")


def build_master_table(metrics: dict[str, float], combined: pd.DataFrame) -> pd.DataFrame:
    combined_accuracy = float(combined["combined_final_correct"].mean())
    common_subset = borrower_subset_label(combined["borrower_id"])
    rows = [
        {
            "module": "Baseline - RouterManager",
            "profiles_run": int(metrics["baseline_profiles"]),
            "disagreement_rate": metrics["baseline_disagreement_rate"],
            "redundancy_index": metrics["raw_redundancy_index"],
            "mean_drift_score": "",
            "final_accuracy": metrics["baseline_accuracy"],
            "key_finding": "Frozen no-intervention benchmark on common subset.",
        },
        {
            "module": "Disagreement Module - Borda",
            "profiles_run": int(metrics["disagreement_profiles"]),
            "disagreement_rate": metrics["disagreement_rate"],
            "redundancy_index": "",
            "mean_drift_score": "",
            "final_accuracy": metrics["borda_accuracy"],
            "key_finding": "Pre-commit, discussion, Borda, CONSENSAGENT, and sycophancy detection active.",
        },
        {
            "module": "Overlap Module - LCE+ToM",
            "profiles_run": int(metrics["overlap_profiles"]),
            "disagreement_rate": "",
            "redundancy_index": metrics["controlled_redundancy_index"],
            "mean_drift_score": "",
            "final_accuracy": metrics["overlap_accuracy"],
            "key_finding": f"Redundancy reduced by {metrics['redundancy_reduction_percent']:.1f}% vs raw overlap.",
        },
        {
            "module": "Drift Module - EMC+ABA",
            "profiles_run": int(metrics["drift_profiles_available"]),
            "disagreement_rate": "",
            "redundancy_index": "",
            "mean_drift_score": metrics["with_aba_mean_drift"],
            "final_accuracy": metrics["drift_router_accuracy_common"],
            "key_finding": f"ABA reduced EMC-only drift by {metrics['aba_reduction_vs_emc_percent']:.1f}%.",
        },
        {
            "module": "Interaction Module - Scenario 4",
            "profiles_run": int(metrics["interaction_profiles"]),
            "disagreement_rate": metrics["co_occurrence_rate"],
            "redundancy_index": metrics["avg_remaining_invocations_after_lce"],
            "mean_drift_score": "",
            "final_accuracy": "",
            "key_finding": "Causal chain measured: overlap, disagreement, resolution, renewed overlap, drift risk.",
        },
        {
            "module": "Scenario 5 Combined Pass",
            "profiles_run": int(combined["borrower_id"].nunique()),
            "disagreement_rate": metrics["co_occurrence_rate"],
            "redundancy_index": metrics["controlled_redundancy_index"],
            "mean_drift_score": metrics["with_aba_mean_drift"],
            "final_accuracy": combined_accuracy,
            "key_finding": f"Integrated Week 4 pass on common {common_subset} subset.",
        },
    ]
    return pd.DataFrame(rows)


def save_rl_rewards(combined: pd.DataFrame) -> pd.DataFrame:
    rewards = combined[
        [
            "run_id",
            "borrower_id",
            "kaggle_truth_label",
            "combined_final_decision",
            "combined_decision_source",
            "combined_reward",
        ]
    ].rename(
        columns={
            "combined_final_decision": "agent_decision",
            "combined_decision_source": "decision_source",
            "combined_reward": "reward",
        }
    )
    rewards.to_csv(RL_REWARDS_PATH, index=False)
    return rewards


def markdown_table(df: pd.DataFrame) -> str:
    return df.to_markdown(index=False)


def plot_figure_5(master: pd.DataFrame, metrics: dict[str, float], combined: pd.DataFrame) -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)

    accuracy_values = {
        "Baseline": metrics["baseline_accuracy"],
        "Disagreement": metrics["borda_accuracy"],
        "Overlap": metrics["overlap_accuracy"],
        "Drift": metrics["drift_router_accuracy_common"],
        "Combined": float(combined["combined_final_correct"].mean()),
    }

    friction_values = {
        "Baseline agent disagreement": metrics["baseline_disagreement_rate"],
        "Disagreement module": metrics["disagreement_rate"],
        "Scenario 4 co-occurrence": metrics["co_occurrence_rate"],
    }

    redundancy_values = {
        "Without LCE": metrics["raw_redundancy_index"],
        "With LCE+ToM": metrics["controlled_redundancy_index"],
    }

    drift_values = {
        "EMC only": metrics["with_emc_mean_drift"],
        "EMC+ABA": metrics["with_aba_mean_drift"],
    }

    fig, axes = plt.subplots(2, 2, figsize=(15, 9.5))
    fig.suptitle("Figure 5 - Week 4 Master Comparison", fontsize=16)

    ax = axes[0, 0]
    ax.bar(accuracy_values.keys(), accuracy_values.values(), color="#2E8B57")
    ax.set_title("A. Final Accuracy")
    ax.set_ylim(0, 1)
    ax.set_ylabel("Accuracy")
    ax.grid(axis="y", alpha=0.25)
    for index, value in enumerate(accuracy_values.values()):
        ax.text(index, value + 0.03, f"{value:.0%}", ha="center")

    ax = axes[0, 1]
    labels = ["Baseline\ndisagreement", "Disagreement\nmodule", "Scenario 4\nco-occurrence"]
    values = list(friction_values.values())
    ax.barh(labels, values, color="#B84A4A")
    ax.set_title("B. Disagreement and Co-occurrence")
    ax.set_xlim(0, 1.08)
    ax.set_xlabel("Rate")
    ax.grid(axis="x", alpha=0.25)
    for index, value in enumerate(values):
        ax.text(value + 0.02, index, f"{value:.0%}", va="center")

    ax = axes[1, 0]
    ax.bar(redundancy_values.keys(), redundancy_values.values(), color=["#9C5E1A", "#D9A441"])
    ax.set_title("C. Redundancy Index")
    ax.set_ylabel("Mean redundancy index")
    ax.grid(axis="y", alpha=0.25)
    for index, value in enumerate(redundancy_values.values()):
        ax.text(index, value + 0.04, f"{value:.2f}", ha="center")

    ax = axes[1, 1]
    ax.bar(drift_values.keys(), drift_values.values(), color=["#2E6F95", "#3A9D73"])
    ax.axhline(0.15, linestyle="--", color="#CC3333", linewidth=1.2, label="KL threshold")
    ax.set_title("D. Mean Drift Score")
    ax.set_ylabel("Mean KL drift")
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    for index, value in enumerate(drift_values.values()):
        ax.text(index, value + 0.03, f"{value:.2f}", ha="center")

    plt.tight_layout(rect=(0, 0, 1, 0.95))
    plt.savefig(FIGURE_PATH, dpi=220)
    plt.close()


def write_report(master: pd.DataFrame, metrics: dict[str, float], combined: pd.DataFrame) -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    common_subset = borrower_subset_label(combined["borrower_id"])
    report = f"""# Week 4 Day 4: Scenario 5 Combined Pipeline Report

## Objective

Thursday integrates the completed Week 3 and Week 4 modules into a single evaluation layer. The integration combines the baseline, disagreement module, overlap module, drift module, and Scenario 4 interaction analysis.

## Scope

This is an integration pass on the common `{common_subset}` subset shared across all available module outputs. Larger 100/200-profile full combined runs are deferred to Week 5 after module-level behaviour is frozen.

## Outputs

- Combined results: `{COMBINED_RESULTS_PATH}`
- Master comparison table: `{MASTER_TABLE_CSV_PATH}`
- RL reward stub: `{RL_REWARDS_PATH}`
- Figure 5: `{FIGURE_PATH}`

## Master Comparison Table

{markdown_table(master)}

## Main Findings

1. Baseline RouterManager accuracy on the common subset was `{metrics['baseline_accuracy']:.1%}`.

2. The overlap module reduced redundancy from `{metrics['raw_redundancy_index']:.2f}` to `{metrics['controlled_redundancy_index']:.2f}`, a `{metrics['redundancy_reduction_percent']:.1f}%` reduction.

3. The drift module reduced EMC-only drift by `{metrics['aba_reduction_vs_emc_percent']:.1f}%` when ABA was added.

4. The integrated Scenario 5 pass achieved `{combined['combined_final_correct'].mean():.1%}` final accuracy on the common subset using the available module decisions.

## RL Reward Stub

The reward signal is intentionally simple at this stage: correct approve/reject decisions receive `+1`, incorrect approve/reject decisions receive `-1`, and abstentions such as refer/unknown receive `0`. This prepares the Week 5 reinforcement learning loop without closing the loop yet.
"""
    REPORT_MD_PATH.write_text(report, encoding="utf-8")

    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading("Week 4 Day 4: Scenario 5 Combined Pipeline Report", level=1)
        doc.add_heading("Objective", level=2)
        doc.add_paragraph(
            "Thursday integrates the completed Week 3 and Week 4 modules into a single evaluation layer: "
            "baseline, disagreement, overlap, drift, and interaction analysis."
        )
        doc.add_heading("Scope", level=2)
        doc.add_paragraph(
            f"This is an integration pass on the common {common_subset} subset shared across all available module outputs. "
            "Larger 100/200-profile combined runs are deferred to Week 5 after module-level behaviour is frozen."
        )
        doc.add_heading("Figure 5 - Master Comparison", level=2)
        doc.add_picture(str(FIGURE_PATH), width=Inches(6.6))
        doc.add_heading("Key Numbers", level=2)
        for bullet in [
            f"Baseline RouterManager accuracy: {metrics['baseline_accuracy']:.1%}",
            f"Disagreement module Borda accuracy: {metrics['borda_accuracy']:.1%}",
            f"Overlap redundancy reduction: {metrics['redundancy_reduction_percent']:.1f}%",
            f"ABA drift reduction vs EMC-only: {metrics['aba_reduction_vs_emc_percent']:.1f}%",
            f"Combined pass final accuracy: {combined['combined_final_correct'].mean():.1%}",
        ]:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_heading("Master Comparison Table", level=2)
        table = doc.add_table(rows=1, cols=len(master.columns))
        table.style = "Table Grid"
        for index, column in enumerate(master.columns):
            table.rows[0].cells[index].text = str(column)
        for _, row in master.iterrows():
            cells = table.add_row().cells
            for index, column in enumerate(master.columns):
                value = row[column]
                if isinstance(value, float):
                    cells[index].text = f"{value:.4f}"
                else:
                    cells[index].text = str(value)
        doc.add_heading("RL Reward Stub", level=2)
        doc.add_paragraph(
            "Correct approve/reject decisions receive +1, incorrect approve/reject decisions receive -1, "
            "and abstentions such as refer/unknown receive 0. This is a logging stub for Week 5."
        )
        doc.save(REPORT_DOCX_PATH)
    except Exception as exc:
        print(f"Could not create DOCX report: {exc}")


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)

    baseline = read_required_csv(BASELINE_PATH)
    disagreement = read_required_csv(DISAGREEMENT_PATH)
    overlap = read_required_csv(OVERLAP_PATH)
    drift_with_emc = read_required_csv(DRIFT_WITH_EMC_PATH)
    drift_with_aba = read_required_csv(DRIFT_WITH_ABA_PATH)
    drift_events = read_required_csv(DRIFT_ABA_EVENTS_PATH)
    dual = read_required_csv(DUAL_RESULTS_PATH)

    borrowers = common_borrowers(baseline, disagreement, overlap, drift_events, dual)
    if not borrowers:
        raise RuntimeError("No common borrowers found across module outputs.")

    baseline_part, baseline_metrics = baseline_table(baseline, borrowers)
    disagreement_part, disagreement_metrics = disagreement_table(disagreement, borrowers)
    overlap_part, overlap_metrics = overlap_table(overlap, borrowers)
    drift_part, drift_metrics = drift_table(drift_events, drift_with_emc, drift_with_aba, borrowers)
    interaction_part, interaction_metrics = interaction_table(dual, borrowers)

    combined = build_combined_results(
        baseline_part,
        disagreement_part,
        overlap_part,
        drift_part,
        interaction_part,
    )

    metrics = {
        **baseline_metrics,
        **disagreement_metrics,
        **overlap_metrics,
        **drift_metrics,
        **interaction_metrics,
    }

    master = build_master_table(metrics, combined)
    rewards = save_rl_rewards(combined)

    combined.to_csv(COMBINED_RESULTS_PATH, index=False)
    master.to_csv(MASTER_TABLE_CSV_PATH, index=False)
    MASTER_TABLE_MD_PATH.write_text(markdown_table(master), encoding="utf-8")
    plot_figure_5(master, metrics, combined)
    write_report(master, metrics, combined)

    print("Scenario 5 combined integration complete.")
    print(f"Common borrowers: {len(borrowers)} ({borrowers[0]}-{borrowers[-1]})")
    print(f"Baseline accuracy: {metrics['baseline_accuracy']:.1%}")
    print(f"Borda accuracy: {metrics['borda_accuracy']:.1%}")
    print(f"Overlap redundancy reduction: {metrics['redundancy_reduction_percent']:.1f}%")
    print(f"ABA drift reduction vs EMC-only: {metrics['aba_reduction_vs_emc_percent']:.1f}%")
    print(f"Combined final accuracy: {combined['combined_final_correct'].mean():.1%}")
    print(f"Rewards logged: {len(rewards)}")
    print(f"Saved: {COMBINED_RESULTS_PATH}")
    print(f"Saved: {MASTER_TABLE_CSV_PATH}")
    print(f"Saved: {RL_REWARDS_PATH}")
    print(f"Saved: {FIGURE_PATH}")
    print(f"Saved: {REPORT_MD_PATH}")
    if REPORT_DOCX_PATH.exists():
        print(f"Saved: {REPORT_DOCX_PATH}")


if __name__ == "__main__":
    main()
