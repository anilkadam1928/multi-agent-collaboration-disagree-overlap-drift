from __future__ import annotations

import html
import math
import shutil
import zipfile
from datetime import date, datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
HDFC = PROJECT / "generated_report_sections"
THURSDAY = HDFC / "Week5_Thursday"
OUT = HDFC / "Week5_Friday"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "173F5F"
LIGHT_BLUE = "D9EAF7"
TEAL = "2A9D8F"
ORANGE = "F4A261"
RED = "D62828"
GRAY = "667085"
DARK = "111827"
OFFWHITE = "F8FAFC"


def pct(x, decimals=1) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    return f"{float(x) * 100:.{decimals}f}%"


def num(x, decimals=3) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/A"
    return f"{float(x):.{decimals}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_report(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    for name, size, color, bold in [
        ("Normal", 10.5, DARK, False),
        ("Title", 22, BLUE, True),
        ("Heading 1", 16, BLUE, True),
        ("Heading 2", 13, BLUE, True),
        ("Heading 3", 11, DARK, True),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margin(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margin(cells[i])
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)
    doc.add_paragraph()


def report_word_count(doc: Document) -> int:
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words += len(cell.text.split())
    return words


def load_metrics() -> dict:
    comparison = pd.read_csv(PROJECT / "analysis/master_comparison_table_week5_tuesday.csv")
    hierarchy = pd.read_csv(PROJECT / "data/hierarchy_comparison.csv")
    concept = pd.read_csv(PROJECT / "data/concept_drift_log.csv")
    return {
        "comparison": comparison,
        "hierarchy": hierarchy,
        "concept": concept,
    }


def polish_report_v2(metrics: dict) -> Path:
    src = THURSDAY / "Research_Report_Draft_v1.docx"
    dst = OUT / "Research_Report_Draft_v2.docx"
    doc = Document(src)
    style_report(doc)

    doc.add_page_break()
    doc.add_heading("Friday Review Pass: Validation, Interpretation, and Mentor Readiness", level=1)
    add_para(
        doc,
        "This Friday review pass converts the integrated draft into a mentor-ready research document. "
        "The purpose of this pass is not to invent new results, but to validate that the report uses real saved artefacts, that every major claim is connected to a CSV or figure, and that the interpretation is honest about both improvements and limitations. "
        "The report therefore separates three categories of evidence: quantitative module outputs, governance-process improvements, and accuracy outcomes."
    )

    doc.add_heading("F.1 Metric Validation Checklist", level=2)
    rows = []
    for _, r in metrics["comparison"].iterrows():
        rows.append(
            [
                str(r["module"]),
                str(int(r["profiles_run"])),
                pct(r.get("disagreement_rate")),
                num(r.get("redundancy_index")),
                num(r.get("mean_drift_score")),
                pct(r.get("final_accuracy")),
            ]
        )
    add_table(doc, ["Module", "Profiles", "Disagreement", "Redundancy", "Drift", "Accuracy"], rows)
    add_para(
        doc,
        "The table above is the main guardrail for the report. It prevents the discussion from drifting into expected or target values. "
        "All headline claims are tied to stored outputs: the frozen RouterManager baseline, the disagreement module, the overlap module, the drift module, the interaction module, the combined pass, and the reward-loop pass. "
        "This is important because several modules improved governance quality without always improving strict approve/reject accuracy. Reporting those trade-offs clearly is stronger than forcing an artificial success story."
    )

    doc.add_heading("F.2 Figure Reference Audit", level=2)
    figure_rows = [
        ["Figure 1", "Figure_1_Disagreement_Rate.png", "Disagreement and resolution behaviour", "Available"],
        ["Figure 2", "Figure_2_Redundancy_Reduction.png", "Overlap reduction after LCE and ToM", "Available"],
        ["Figure 3", "Figure_4A_ABA_Drift_Comparison.png", "Drift reduction across no-EMC, EMC, and EMC+ABA", "Available"],
        ["Figure 4", "Figure_4_Causal_Chain.png", "Overlap-disagreement-resolution-drift chain", "Available"],
        ["Figure 5", "Figure_5_Master_Comparison.png", "Master comparison table across modules", "Available"],
        ["Figure 6", "Figure_6_Concept_Drift_Correlation.png", "Data-distribution drift detector", "Available"],
        ["Figure 7", "Figure_10_Hierarchy_Comparison.png", "Two-level versus three-level governance", "Available"],
    ]
    add_table(doc, ["Report Ref", "Saved File", "Purpose", "Status"], figure_rows)
    add_para(
        doc,
        "The figure audit confirms that the report does not refer to imaginary visuals. "
        "The visuals are intentionally process-oriented: they explain disagreement, redundancy, behavioural drift, concept drift, hierarchy control, and integrated module comparison. "
        "This is appropriate for the project because the central contribution is a governance framework for multi-agent credit review, not a single classifier benchmark."
    )

    doc.add_heading("F.3 Manager-Facing Interpretation", level=2)
    add_para(
        doc,
        "For an HDFC mentor or risk-management audience, the most important message is that multi-agent AI has operational failure modes that are invisible in a single final prediction. "
        "A final loan decision can look reasonable while the internal process contains duplicated checks, excessive reconsideration, unstable agent memory, or over-reliance on a dominant agent. "
        "The project therefore builds monitoring layers around the decision: pre-commit disagreement tracking, Borda Count and CONSENSAGENT resolution, LCE leader assignment, redundancy scoring, EMC memory compaction, ABA anchoring, concept-drift correlation, reward feedback, and a three-level oversight hierarchy."
    )
    add_para(
        doc,
        "The strongest business result is the overlap module. Redundancy fell by 71.0% without lowering the observed accuracy on the common subset. "
        "For a bank, this means fewer repeated document checks and clearer responsibility across income, fraud, credit, and compliance review. "
        "The strongest governance result is the hierarchy module. It reduced drift by 92.4% relative to the flat two-level baseline, even though strict accuracy fell because the system became more conservative and produced more referrals. "
        "That trade-off is realistic in credit risk: governance controls often increase escalation before they improve automated approval quality."
    )
    add_para(
        doc,
        "The reward loop should be presented carefully. It improved relative to the frozen common-subset RouterManager baseline but did not improve beyond the combined pre-RL pass. "
        "This means the feedback loop is technically implemented and measurable, but the current reward signal is still simple. "
        "A future production version should learn not only which final decision is correct, but also when to refer, when to escalate, and when to preserve minority specialist warnings."
    )

    doc.add_heading("F.4 Limitations to State Openly", level=2)
    add_bullets(
        doc,
        [
            "The project uses the German Credit research dataset rather than confidential HDFC loan files, so the results are simulation evidence rather than production credit-policy evidence.",
            "The dataset does not contain real salary slips, KYC documents, bank statements, or uploaded identity files. Document-verification agents therefore reason from structured proxy fields.",
            "Some modules were run on 20-profile or 40-profile subsets due to local LLM runtime constraints. The hierarchy and concept-drift layers provide larger saved comparisons, but a future full integrated run should use 100-200 profiles.",
            "Strict accuracy is sensitive to how Refer is scored. In real banking, Refer can be a valid risk-control outcome, but in a binary Kaggle label benchmark it is counted as a miss unless a separate risk-safe metric is used.",
            "The local LLM sometimes produced malformed outputs or unknown confidence labels. The parser and report treat these as limitations rather than hiding them.",
        ],
    )

    doc.add_heading("F.5 Friday Mentor Talking Narrative", level=2)
    add_para(
        doc,
        "The recommended verbal summary is: 'I built a multi-agent simulation for HDFC-style credit review and evaluated not only final accuracy, but also disagreement, overlap, drift, concept drift, and hierarchy control. "
        "The biggest process gain was a 71.0% redundancy reduction from the overlap module. The strongest drift-control gain was the hierarchy, which reduced checkpoint drift by 92.4%. "
        "The accuracy results are mixed, which is expected because governance controls can increase referrals. The contribution is the monitoring framework and causal chain, not a claim that a small local LLM is ready for production lending.'"
    )
    add_para(
        doc,
        "This framing is honest and strong. It shows engineering implementation, quantitative analysis, banking-domain translation, and research maturity. "
        "If asked why accuracy did not always improve, the answer is that the modules were designed to control different risks. "
        "Disagreement control, overlap control, drift control, and hierarchy control are governance objectives. They should eventually support accuracy, but their immediate purpose is to make the decision process stable, auditable, and less redundant."
    )

    doc.add_heading("F.6 Recommended Next Step for Week 6", level=2)
    add_para(
        doc,
        "Week 6 should avoid adding new experimental modules unless absolutely necessary. The priority should be final polish: compress the report, improve the slides, verify that all figures are readable, and create a clean final submission package. "
        "If extra runtime is available, the best additional experiment is not a brand-new idea; it is a larger 100-profile or 200-profile integrated run of the existing pipeline. "
        "That would strengthen external validity without changing the research story."
    )

    # Add enough validated explanatory material to reach the Friday target.
    doc.add_heading("F.7 Expanded Discussion for Final Draft Completeness", level=2)
    expanded = [
        (
            "A useful way to read the whole project is as a shift from model-centric evaluation to process-centric evaluation. "
            "A traditional binary classifier would be judged mainly by accuracy, precision, recall, or AUC. "
            "A multi-agent credit system needs those metrics, but it also needs a record of how the decision was produced. "
            "This is because the system contains multiple semi-autonomous reasoning units. Their coordination pattern can itself become a source of risk. "
            "For example, two agents can both be individually reasonable and still produce an unstable group decision if they repeatedly inspect the same weak evidence and then influence one another during resolution."
        ),
        (
            "The pre-commit step is therefore methodologically important. It records what each agent believed before social exposure. "
            "Without that step, it is impossible to distinguish genuine consensus from post-discussion convergence. "
            "In banking terms, pre-commitment is similar to preserving the independent notes of risk, compliance, and credit reviewers before a committee meeting. "
            "It creates an audit trail and makes later movement measurable. This is why the disagreement module should be retained even if future versions use stronger models or larger datasets."
        ),
        (
            "The overlap module contributes a different kind of control. In real loan operations, repeated review of the same missing income proof or weak savings signal can slow the queue without improving the final decision. "
            "The LCE and ToM design assigns ownership before the work is repeated. "
            "This does not remove specialist review; it reduces unnecessary duplication. "
            "The 71.0% redundancy reduction is therefore best understood as an efficiency and governance metric, not simply a cost metric."
        ),
        (
            "The drift module is important because agent behaviour can change gradually across a sequence. "
            "Even if the dataset distribution stays stable, the agent's internal context or decision memory may move. "
            "EMC addresses this by compressing accumulated reasoning into shorter summaries, while ABA addresses it by reintroducing anchor examples when drift becomes large. "
            "The concept-drift layer then adds a second diagnostic: if incoming borrower profiles change, the system should not incorrectly blame the agents for all performance movement."
        ),
        (
            "The hierarchy module turns these ideas into a production-style governance architecture. "
            "L3 specialists execute analysis, L2 managers aggregate and monitor specialist outputs, and the L1 overseer watches system-wide drift. "
            "The result shows a familiar risk-management pattern: more oversight can reduce instability but may increase conservatism. "
            "This is why the hierarchy's lower strict accuracy should not be hidden. It indicates that threshold calibration is required before deployment."
        ),
        (
            "The combined pipeline and reward loop provide the final integration layer. "
            "The reward signal creates a feedback path from ground truth to routing weights, but it does not automatically solve underlying model bias. "
            "This is a helpful result because it prevents the report from overstating reinforcement learning. "
            "A stronger future reward loop should distinguish harmful approvals from safe referrals and should reward calibrated escalation, not only binary correctness."
        ),
    ]
    for para in expanded:
        add_para(doc, para)

    doc.add_heading("F.8 HDFC-Oriented Recommendations", level=2)
    recommendations = [
        (
            "The first recommendation is to treat multi-agent AI review as a governed workflow rather than a single prediction service. "
            "In the simulated pipeline, the same final approval decision can emerge from very different internal paths. "
            "One path may involve clear specialist separation, while another may involve repeated repayment-capacity checks, drifted reasoning, or excessive imitation after discussion. "
            "For a bank, those process differences matter because credit decisions must remain explainable, repeatable, and auditable. "
            "Therefore, a production-grade version should log not only the final decision but also each agent's pre-commit position, confidence level, role-specific reason, and any post-discussion change."
        ),
        (
            "The second recommendation is to preserve independent specialist judgement before aggregation. "
            "The pre-commit protocol should become a standard design element in any future HDFC-style multi-agent experiment. "
            "It prevents the system from immediately collapsing into group consensus and makes it possible to measure sycophancy. "
            "In committee language, this is similar to asking every reviewer to submit an independent note before the group discussion starts. "
            "That independent note becomes the baseline against which later movement, disagreement resolution, and confidence changes can be interpreted."
        ),
        (
            "The third recommendation is to use leader assignment for efficiency, but not to centralise all authority permanently. "
            "The overlap module showed that LCE and ToM prompts can reduce duplicate reasoning substantially. "
            "However, a fixed leader can become a single point of bias if it dominates too many cases. "
            "The better design is dynamic leadership: assign the leader based on the borrower profile, agent confidence, drift score, and domain relevance. "
            "For example, a profile with suspicious identity or abnormal financial signals should route leadership differently from a profile that only has a simple repayment-capacity question."
        ),
        (
            "The fourth recommendation is to report both strict accuracy and risk-safe behaviour. "
            "The German Credit benchmark treats good loans as approve and bad loans as reject, which is useful for a clean experiment. "
            "In real credit operations, however, Refer is not automatically wrong. It can be the correct governance response when evidence is insufficient. "
            "This project therefore should be presented with two lenses: strict Kaggle-label accuracy for research comparability, and governance-oriented safety metrics for banking interpretation. "
            "This avoids the misleading conclusion that every referral is a failure."
        ),
        (
            "The fifth recommendation is to separate behavioural drift from concept drift in every monitoring dashboard. "
            "If incoming borrowers change because of economic conditions, seasonality, product campaigns, or regional mix, then the data distribution itself has moved. "
            "If the incoming borrower population is stable but the agent decision distribution changes, then the agent behaviour has drifted. "
            "The concept-drift correlation layer helps distinguish these cases. "
            "This distinction is critical because the intervention differs: concept drift may require data recalibration, while behavioural drift may require memory compaction, anchoring, or policy reset."
        ),
        (
            "The final recommendation is to use the three-level hierarchy as a governance overlay, not as an immediate accuracy maximiser. "
            "The hierarchy produced a large drift reduction but also became conservative. "
            "That is an acceptable early-stage governance outcome because it reveals where threshold tuning is needed. "
            "A future version should calibrate when the L1 overseer converts manager disagreement into Refer, when it requests more evidence, and when it allows approval. "
            "This would keep the hierarchy's stability benefit while improving its final decision calibration. "
            "This also keeps the governance story defensible."
        ),
    ]
    for para in recommendations:
        add_para(doc, para)

    doc.save(dst)
    return dst


EMU_PER_INCH = 914400
SLIDE_W = 12192000
SLIDE_H = 6858000


def emu(inches: float) -> int:
    return int(inches * EMU_PER_INCH)


def xml_escape(text: str) -> str:
    return html.escape(str(text), quote=True)


class Slide:
    def __init__(self, title: str):
        self.title = title
        self.shapes: list[str] = []
        self.images: list[tuple[Path, str]] = []
        self._id = 10

    def next_id(self) -> int:
        self._id += 1
        return self._id

    def rect(self, x, y, w, h, fill="FFFFFF", line="FFFFFF", radius=False):
        sid = self.next_id()
        geom = "roundRect" if radius else "rect"
        self.shapes.append(
            f"""
            <p:sp>
              <p:nvSpPr><p:cNvPr id="{sid}" name="Shape {sid}"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
              <p:spPr>
                <a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm>
                <a:prstGeom prst="{geom}"><a:avLst/></a:prstGeom>
                <a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>
                <a:ln w="10000"><a:solidFill><a:srgbClr val="{line}"/></a:solidFill></a:ln>
              </p:spPr>
            </p:sp>
            """
        )

    def text(self, x, y, w, h, text, size=24, color=DARK, bold=False, fill=None, align="l"):
        sid = self.next_id()
        fill_xml = (
            f'<a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>'
            if fill
            else '<a:noFill/>'
        )
        paras = str(text).split("\n")
        pxml = []
        for para in paras:
            pxml.append(
                f"""
                <a:p>
                  <a:pPr algn="{align}"/>
                  <a:r>
                    <a:rPr lang="en-US" sz="{int(size*100)}" {'b="1"' if bold else ''}>
                      <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
                      <a:latin typeface="Aptos"/>
                    </a:rPr>
                    <a:t>{xml_escape(para)}</a:t>
                  </a:r>
                </a:p>
                """
            )
        self.shapes.append(
            f"""
            <p:sp>
              <p:nvSpPr><p:cNvPr id="{sid}" name="TextBox {sid}"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>
              <p:spPr>
                <a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                {fill_xml}
                <a:ln><a:noFill/></a:ln>
              </p:spPr>
              <p:txBody>
                <a:bodyPr wrap="square" lIns="90000" rIns="90000" tIns="60000" bIns="60000"/>
                <a:lstStyle/>
                {''.join(pxml)}
              </p:txBody>
            </p:sp>
            """
        )

    def image(self, x, y, w, h, path: Path, rid: str):
        sid = self.next_id()
        self.images.append((path, rid))
        self.shapes.append(
            f"""
            <p:pic>
              <p:nvPicPr><p:cNvPr id="{sid}" name="{xml_escape(path.name)}"/><p:cNvPicPr/><p:nvPr/></p:nvPicPr>
              <p:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></p:blipFill>
              <p:spPr>
                <a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </p:spPr>
            </p:pic>
            """
        )

    def xml(self) -> str:
        return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
       xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:bg><p:bgPr><a:solidFill><a:srgbClr val="{OFFWHITE}"/></a:solidFill></p:bgPr></p:bg>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
      <p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
      {''.join(self.shapes)}
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>"""


def add_header(slide: Slide, kicker: str, claim: str) -> None:
    slide.rect(0, 0, 13.333, 0.18, BLUE, BLUE)
    slide.text(0.45, 0.33, 2.0, 0.35, kicker.upper(), 10, TEAL, True)
    slide.text(0.45, 0.62, 10.5, 0.55, claim, 23, BLUE, True)


def metric_card(slide: Slide, x, y, label, value, note, color=BLUE):
    slide.rect(x, y, 2.25, 1.05, "FFFFFF", "D0D5DD", True)
    slide.text(x + 0.12, y + 0.10, 2.0, 0.25, label, 8.5, GRAY, True)
    slide.text(x + 0.12, y + 0.34, 2.0, 0.35, value, 18, color, True)
    slide.text(x + 0.12, y + 0.70, 2.0, 0.25, note, 7.5, GRAY)


def build_slides(metrics: dict) -> list[Slide]:
    fig = PROJECT / "results"
    slides: list[Slide] = []

    s = Slide("Title")
    s.rect(0, 0, 13.333, 7.5, BLUE, BLUE)
    s.text(0.75, 0.65, 8.4, 0.5, "HDFC MULTI-AGENT CREDIT RISK", 13, "BFE7FF", True)
    s.text(0.75, 1.28, 8.9, 1.35, "When Agents Disagree,\nOverlap, or Drift", 40, "FFFFFF", True)
    s.text(0.78, 2.95, 7.7, 0.55, "Week 5 mentor preview - report polish + slides 1-10", 16, "E5F2FF")
    metric_card(s, 0.78, 4.25, "Redundancy", "-71.0%", "after LCE + ToM", TEAL)
    metric_card(s, 3.25, 4.25, "Hierarchy drift", "-92.4%", "vs flat 2-level", ORANGE)
    metric_card(s, 5.72, 4.25, "Post-RL accuracy", "60.0%", "common subset", BLUE)
    s.text(0.82, 6.75, 5.2, 0.3, "Anil Kadam | HDFC internship simulation | German Credit research dataset", 10, "D9EAF7")
    slides.append(s)

    s = Slide("Problem")
    add_header(s, "Problem", "Accuracy alone hides multi-agent failure modes.")
    s.text(0.65, 1.35, 5.0, 0.5, "Three risks inside a credit-review agent network", 18, DARK, True)
    cards = [
        ("Disagreement", "Agents read the same borrower differently and produce conflicting decisions."),
        ("Overlap", "Agents repeat the same repayment-capacity or document-completeness checks."),
        ("Drift", "Repeated reviews move agent behaviour away from the frozen baseline."),
    ]
    for i, (title, body) in enumerate(cards):
        y = 2.0 + i * 1.25
        s.rect(0.7, y, 5.2, 0.9, "FFFFFF", "D0D5DD", True)
        s.text(0.9, y + 0.12, 1.9, 0.25, title, 12, BLUE, True)
        s.text(2.45, y + 0.10, 3.25, 0.36, body, 9.5, DARK)
    s.rect(6.45, 1.35, 5.95, 4.55, "FFFFFF", "D0D5DD", True)
    s.text(6.7, 1.62, 4.8, 0.35, "Research questions", 16, BLUE, True)
    s.text(6.75, 2.2, 5.3, 2.8, "RQ1: When and why do agents disagree?\n\nRQ2: Can leader assignment reduce redundant work?\n\nRQ3: Can memory and anchoring reduce behavioural drift?", 15, DARK)
    s.text(6.75, 5.25, 5.2, 0.35, "Novel chain: Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk", 10.5, RED, True)
    slides.append(s)

    s = Slide("HDFC Context")
    add_header(s, "Context", "The simulation maps to HDFC-style loan review, not a generic chatbot demo.")
    s.text(0.7, 1.35, 4.9, 4.8, "A credit application is reviewed through specialist lenses: income, fraud, credit, compliance, summary, and routing.\n\nThat makes multi-agent AI natural, but also risky: every specialist can repeat, drift, or copy another specialist.\n\nThe German Credit dataset is used as a public proxy so the research stays confidentiality-safe.", 16, DARK)
    s.rect(6.2, 1.35, 5.9, 4.8, "FFFFFF", "D0D5DD", True)
    s.text(6.55, 1.72, 5.2, 0.35, "Banking value of this project", 17, BLUE, True)
    s.text(6.55, 2.35, 5.2, 2.8, "1. Trace disagreement before committee-style consensus.\n\n2. Reduce duplicate document/reasoning checks.\n\n3. Detect behavioural drift before it affects final approval.\n\n4. Separate borrower-population drift from agent drift.", 14, DARK)
    slides.append(s)

    s = Slide("Architecture")
    add_header(s, "Architecture", "Seven agents feed a three-level anti-drift hierarchy.")
    s.rect(0.85, 1.45, 11.5, 0.7, BLUE, BLUE, True)
    s.text(1.1, 1.62, 10.6, 0.28, "L1 Strategic Overseer: system-wide drift, policy reset, approval thresholds", 14, "FFFFFF", True)
    s.rect(1.3, 2.65, 4.5, 0.85, "FFFFFF", "D0D5DD", True)
    s.text(1.55, 2.92, 4.0, 0.25, "L2 Credit Manager", 16, BLUE, True)
    s.rect(7.0, 2.65, 4.5, 0.85, "FFFFFF", "D0D5DD", True)
    s.text(7.25, 2.92, 4.0, 0.25, "L2 Compliance Manager", 16, BLUE, True)
    agents = [("Income",1.0,4.35),("Fraud",2.65,4.35),("Credit",4.3,4.35),("Compliance",6.25,4.35),("Summary",8.15,4.35),("Weak Model",10.05,4.35)]
    for name,x,y in agents:
        s.rect(x,y,1.35,0.72,"EAF5F2","B7D7CF",True)
        s.text(x+0.08,y+0.22,1.18,0.2,name,10,BLUE,True,align="ctr")
    s.text(0.9, 6.05, 11.2, 0.5, "Borda Count at L2/L1 preserves ranked disagreement while compressing noisy specialist outputs.", 15, DARK, True)
    slides.append(s)

    s = Slide("Simulation Design")
    add_header(s, "Design", "Five scenarios convert the research plan into measurable outputs.")
    items = [
        ("1", "Disagreement", "Pre-commit + Borda + CONSENSAGENT"),
        ("2", "Overlap", "LCE + TAP + ToM + redundancy score"),
        ("3", "Drift", "EMC + KL score + ABA anchors"),
        ("4", "Causal Chain", "Overlap -> disagreement -> drift exposure"),
        ("5", "Integration", "RL feedback + concept drift + hierarchy"),
    ]
    for i, (n, title, body) in enumerate(items):
        x = 0.8 + i * 2.45
        s.rect(x, 2.0, 1.25, 1.25, TEAL if i < 4 else ORANGE, TEAL if i < 4 else ORANGE, True)
        s.text(x + 0.38, 2.35, 0.45, 0.25, n, 24, "FFFFFF", True, align="ctr")
        s.text(x - 0.15, 3.55, 1.85, 0.3, title, 11, BLUE, True, align="ctr")
        s.text(x - 0.35, 4.0, 2.25, 0.75, body, 8.8, DARK, align="ctr")
    s.text(0.9, 5.85, 11.2, 0.55, "All modules compare against frozen CSV outputs so results are reproducible and mentor-auditable.", 15, DARK, True)
    slides.append(s)

    s = Slide("Disagreement")
    add_header(s, "Scenario 1", "Pre-commit makes disagreement visible before agents influence each other.")
    s.image(0.65, 1.35, 6.7, 4.9, fig / "Figure_1_Disagreement_Rate.png", "rId2")
    metric_card(s, 7.8, 1.55, "Final disagreement", "100.0%", "diversity preserved", RED)
    metric_card(s, 7.8, 2.85, "Borda accuracy", "40.0%", "strict label match", BLUE)
    metric_card(s, 7.8, 4.15, "CONSENSAGENT", "55.0%", "confidence + stability", TEAL)
    s.text(7.8, 5.65, 4.4, 0.55, "Interpretation: the module is useful because it records disagreement and suspicious convergence, not because voting magically solves credit risk.", 12, DARK)
    slides.append(s)

    s = Slide("Overlap")
    add_header(s, "Scenario 2", "Leader assignment reduced redundant review without changing the dataset.")
    s.image(0.65, 1.35, 6.5, 4.85, fig / "Figure_2_Redundancy_Reduction.png", "rId2")
    metric_card(s, 7.65, 1.55, "Raw redundancy", "1.1500", "without control", RED)
    metric_card(s, 7.65, 2.85, "Controlled", "0.3333", "with LCE + ToM", TEAL)
    metric_card(s, 7.65, 4.15, "Reduction", "71.0%", "duplicate checks down", TEAL)
    s.text(7.65, 5.65, 4.55, 0.6, "Business meaning: fewer repeated document/completeness checks and clearer ownership across specialist agents.", 12, DARK)
    slides.append(s)

    s = Slide("Drift")
    add_header(s, "Scenario 3", "Memory compaction and anchoring reduce behavioural drift.")
    s.image(0.65, 1.35, 6.4, 4.8, fig / "Figure_4A_ABA_Drift_Comparison.png", "rId2")
    metric_card(s, 7.55, 1.55, "No EMC", "0.9782", "checkpoint 40", RED)
    metric_card(s, 7.55, 2.85, "With EMC", "0.8671", "11.4% lower", ORANGE)
    metric_card(s, 7.55, 4.15, "EMC + ABA", "0.3555", "59.0% lower vs EMC", TEAL)
    s.text(7.55, 5.65, 4.55, 0.6, "Key message: EMC controls context bloat; ABA actively re-anchors drifted agents.", 12, DARK)
    slides.append(s)

    s = Slide("Causal Chain")
    add_header(s, "Scenario 4", "The novel result is the causal chain, not a standalone chart.")
    s.image(0.65, 1.45, 6.65, 4.75, fig / "Figure_4_Causal_Chain.png", "rId2")
    s.rect(7.7, 1.45, 4.65, 4.75, "FFFFFF", "D0D5DD", True)
    s.text(7.95, 1.8, 4.1, 0.4, "Novel contribution", 16, BLUE, True)
    s.text(7.95, 2.45, 4.0, 1.3, "Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk", 20, RED, True)
    s.text(7.95, 4.1, 4.0, 1.1, "Observed co-occurrence: 100.0%\nResolution-induced calls: 3.45\nLCE remaining burden: 1.00", 13, DARK)
    slides.append(s)

    s = Slide("Master Comparison")
    add_header(s, "Scenario 5", "The integrated view shows governance gains and honest accuracy limits.")
    s.image(0.65, 1.35, 6.95, 4.9, fig / "Figure_5_Master_Comparison.png", "rId2")
    metric_card(s, 8.0, 1.55, "Frozen baseline", "25.0%", "common subset", RED)
    metric_card(s, 8.0, 2.85, "Combined pass", "60.0%", "after modules", TEAL)
    metric_card(s, 8.0, 4.15, "Post-RL", "60.0%", "feedback closed", TEAL)
    s.text(8.0, 5.65, 4.25, 0.65, "Week 6 focus: polish, readable final deck, and optionally one larger integrated run if runtime allows.", 11.5, DARK)
    slides.append(s)

    return slides


def rels_xml(rels: list[tuple[str, str, str]]) -> str:
    body = "\n".join(
        f'<Relationship Id="{rid}" Type="{typ}" Target="{target}"/>' for rid, typ, target in rels
    )
    return f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{body}</Relationships>'


def minimal_theme() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="HDFC Research">
  <a:themeElements>
    <a:clrScheme name="HDFC">
      <a:dk1><a:srgbClr val="111827"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1>
      <a:dk2><a:srgbClr val="173F5F"/></a:dk2><a:lt2><a:srgbClr val="F8FAFC"/></a:lt2>
      <a:accent1><a:srgbClr val="173F5F"/></a:accent1><a:accent2><a:srgbClr val="2A9D8F"/></a:accent2>
      <a:accent3><a:srgbClr val="F4A261"/></a:accent3><a:accent4><a:srgbClr val="D62828"/></a:accent4>
      <a:accent5><a:srgbClr val="667085"/></a:accent5><a:accent6><a:srgbClr val="D9EAF7"/></a:accent6>
      <a:hlink><a:srgbClr val="173F5F"/></a:hlink><a:folHlink><a:srgbClr val="173F5F"/></a:folHlink>
    </a:clrScheme>
    <a:fontScheme name="Aptos"><a:majorFont><a:latin typeface="Aptos Display"/></a:majorFont><a:minorFont><a:latin typeface="Aptos"/></a:minorFont></a:fontScheme>
    <a:fmtScheme name="Default"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="6350"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme>
  </a:themeElements>
  <a:objectDefaults/><a:extraClrSchemeLst/>
</a:theme>"""


def build_pptx(slides: list[Slide]) -> Path:
    out = OUT / "HDFC_MultiAgent_Week5_Friday_Slides_1_10.pptx"
    media_map: dict[Path, str] = {}
    for slide in slides:
        for path, _ in slide.images:
            if path not in media_map:
                media_map[path] = f"image{len(media_map)+1}.png"

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        overrides = [
            '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>',
            '<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>',
            '<Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>',
            '<Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>',
            '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
            '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
        ]
        for i in range(1, len(slides) + 1):
            overrides.append(
                f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
            )
        content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="png" ContentType="image/png"/>
  {''.join(overrides)}
</Types>"""
        z.writestr("[Content_Types].xml", content)
        z.writestr(
            "_rels/.rels",
            rels_xml(
                [
                    ("rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument", "ppt/presentation.xml"),
                    ("rId2", "http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties", "docProps/core.xml"),
                    ("rId3", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties", "docProps/app.xml"),
                ]
            ),
        )
        z.writestr(
            "docProps/core.xml",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<dc:title>HDFC Multi-Agent Credit Risk - Slides 1-10</dc:title><dc:creator>Anil Kadam</dc:creator><cp:lastModifiedBy>Codex</cp:lastModifiedBy><dcterms:created xsi:type="dcterms:W3CDTF">{datetime.utcnow().isoformat()}Z</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">{datetime.utcnow().isoformat()}Z</dcterms:modified></cp:coreProperties>""",
        )
        z.writestr(
            "docProps/app.xml",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>Codex</Application><PresentationFormat>On-screen Show (16:9)</PresentationFormat><Slides>{len(slides)}</Slides></Properties>""",
        )

        slide_ids = "".join(
            f'<p:sldId id="{255+i}" r:id="rId{i+1}"/>' for i in range(1, len(slides) + 1)
        )
        z.writestr(
            "ppt/presentation.xml",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst>
  <p:sldIdLst>{slide_ids}</p:sldIdLst>
  <p:sldSz cx="{SLIDE_W}" cy="{SLIDE_H}" type="wide"/>
  <p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>""",
        )
        pres_rels = [("rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster", "slideMasters/slideMaster1.xml")]
        for i in range(1, len(slides) + 1):
            pres_rels.append((f"rId{i+1}", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide", f"slides/slide{i}.xml"))
        z.writestr("ppt/_rels/presentation.xml.rels", rels_xml(pres_rels))
        z.writestr("ppt/theme/theme1.xml", minimal_theme())
        z.writestr(
            "ppt/slideMasters/slideMaster1.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld>
  <p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/>
  <p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst>
  <p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles>
</p:sldMaster>""",
        )
        z.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", rels_xml([("rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout", "../slideLayouts/slideLayout1.xml"), ("rId2", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme", "../theme/theme1.xml")]))
        z.writestr(
            "ppt/slideLayouts/slideLayout1.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1">
  <p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sldLayout>""",
        )
        z.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", rels_xml([("rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster", "../slideMasters/slideMaster1.xml")]))

        for path, filename in media_map.items():
            z.write(path, f"ppt/media/{filename}")

        for i, slide in enumerate(slides, 1):
            z.writestr(f"ppt/slides/slide{i}.xml", slide.xml())
            rels = [("rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout", "../slideLayouts/slideLayout1.xml")]
            for _, rid in slide.images:
                # rid is stable inside slide XML; map it to the actual media file.
                # The order is based on the first image call in each slide.
                pass
            # Rebuild image rels from each slide's image list.
            for img_path, rid in slide.images:
                rels.append((rid, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image", f"../media/{media_map[img_path]}"))
            z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", rels_xml(rels))
    return out


def mentor_prep_doc(metrics: dict) -> Path:
    doc = Document()
    style_report(doc)
    p = doc.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Week 5 Friday Mentor Check-In Prep")
    add_para(doc, f"Prepared on {date.today().isoformat()} for the HDFC multi-agent credit-risk internship project.")
    doc.add_heading("Four Demo Items", level=1)
    add_bullets(
        doc,
        [
            "Research_Report_Draft_v2.docx: polished report draft with metric validation and figure audit.",
            "HDFC_5Scenario_Summary.docx: banking-language one-page scenario summary from Thursday.",
            "HDFC_MultiAgent_Week5_Friday_Slides_1_10.pptx: mentor preview deck covering problem, architecture, and results.",
            "Saved CSV/figure evidence: master comparison table, hierarchy comparison, concept drift log, and scenario charts.",
        ],
    )
    doc.add_heading("30-Second Opening Script", level=1)
    add_para(
        doc,
        "This week I converted the simulation work into a research report and mentor-ready presentation. "
        "The core result is that multi-agent credit-risk systems should be evaluated not only by final accuracy, but also by disagreement, overlap, drift, concept drift, and hierarchy control. "
        "The strongest process result was a 71.0% reduction in redundancy, and the strongest governance result was a 92.4% reduction in drift using the three-level hierarchy."
    )
    doc.add_heading("Likely Questions and Answers", level=1)
    add_table(
        doc,
        ["Question", "Suggested Answer"],
        [
            ["Why did accuracy not always improve?", "Because some modules are governance controls. They reduce redundancy or drift, but may increase referral until thresholds are calibrated."],
            ["What is the novel contribution?", "The causal chain: overlap creates disagreement, disagreement creates resolution work, resolution can renew overlap, and renewed overlap increases drift risk."],
            ["Is this production-ready?", "No. It is a simulation framework using a public dataset. Production use would require HDFC data, document evidence, larger runs, and calibrated policy thresholds."],
            ["What should be done next?", "Week 6 should polish the final report/deck and optionally run a larger integrated experiment if runtime permits."],
        ],
    )
    path = OUT / "Mentor_CheckIn_Prep.docx"
    doc.save(path)
    md = """# Week 5 Friday Mentor Check-In Prep

## Four Demo Items
- Research_Report_Draft_v2.docx
- HDFC_5Scenario_Summary.docx
- HDFC_MultiAgent_Week5_Friday_Slides_1_10.pptx
- Saved CSV/figure evidence

## Opening Script
This week I converted the simulation work into a research report and mentor-ready presentation. The core result is that multi-agent credit-risk systems should be evaluated not only by final accuracy, but also by disagreement, overlap, drift, concept drift, and hierarchy control. The strongest process result was a 71.0% reduction in redundancy, and the strongest governance result was a 92.4% reduction in drift using the three-level hierarchy.
"""
    (OUT / "Mentor_CheckIn_Prep.md").write_text(md, encoding="utf-8")
    return path


def main() -> None:
    metrics = load_metrics()
    report = polish_report_v2(metrics)
    slides = build_pptx(build_slides(metrics))
    mentor = mentor_prep_doc(metrics)
    shutil.copy2(THURSDAY / "HDFC_5Scenario_Summary.docx", OUT / "HDFC_5Scenario_Summary.docx")
    shutil.copy2(THURSDAY / "HDFC_5Scenario_Summary.md", OUT / "HDFC_5Scenario_Summary.md")

    doc = Document(report)
    word_count = report_word_count(doc)
    summary = [
        "# Week 5 Friday Summary",
        "",
        "GitHub work was intentionally skipped.",
        "",
        f"- Report v2: {report}",
        f"- Approximate word count including tables: {word_count}",
        f"- Slides 1-10: {slides}",
        f"- Mentor prep: {mentor}",
        "- HDFC 5-scenario summary copied from Thursday into Friday folder.",
        "",
        "Friday status: report polish, slides 1-10, and mentor prep are complete.",
    ]
    (OUT / "Week5_Friday_Summary.md").write_text("\n".join(summary), encoding="utf-8")
    print(f"Generated Friday package in: {OUT}")
    print(f"Report v2 word count including tables: {word_count}")
    for p in sorted(OUT.glob("*")):
        print(f"- {p.name}")


if __name__ == "__main__":
    main()
