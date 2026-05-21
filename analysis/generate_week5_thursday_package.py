from __future__ import annotations

import math
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
HDFC_ROOT = PROJECT_ROOT / "generated_report_sections"
OUT_DIR = HDFC_ROOT / "Week5_Thursday"
OUT_DIR.mkdir(parents=True, exist_ok=True)

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"


def pct(value: float | int | None, decimals: int = 1) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "N/A"
    return f"{float(value) * 100:.{decimals}f}%"


def num(value: float | int | None, decimals: int = 3) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "N/A"
    return f"{float(value):.{decimals}f}"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def strip_md_heading(text: str) -> str:
    lines = []
    for line in text.splitlines():
        if line.strip().startswith("#"):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def paragraph_chunks(text: str) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if current:
                chunks.append(" ".join(current).strip())
                current = []
            continue
        if line.startswith("|"):
            continue
        current.append(line)
    if current:
        chunks.append(" ".join(current).strip())
    return [c for c in chunks if c]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_pct: int = 5000) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_w.set(qn("w:type"), "pct")


def style_document(doc: Document, label: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Title", 22, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "000000"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header
    p = header.paragraphs[0]
    p.text = label
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "HDFC multi-agent credit-risk simulation - internal research draft"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.name = "Arial"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_title(doc: Document, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    doc.add_paragraph()
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    for i, (key, value) in enumerate(meta):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_margin(cell)
            for p2 in cell.paragraphs:
                for r in p2.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
            cells[0].paragraphs[0].runs[0].font.bold = True


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Arial"
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_margin(hdr[i])
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.bold = True
                r.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_margin(cells[i])
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margin(cell, 140, 160, 140, 160)
    cell.text = ""
    p = cell.paragraphs[0]
    t = p.add_run(title + ": ")
    t.bold = True
    t.font.name = "Arial"
    t.font.size = Pt(10.5)
    body = p.add_run(text)
    body.font.name = "Arial"
    body.font.size = Pt(10.5)
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.3) -> None:
    if not image_path.exists():
        add_callout(doc, "Missing figure", f"{image_path.name} was not found at generation time.")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)


def load_metrics() -> dict:
    m: dict = {}
    comparison = pd.read_csv(PROJECT_ROOT / "analysis/master_comparison_table_week5_tuesday.csv")
    hierarchy = pd.read_csv(PROJECT_ROOT / "data/hierarchy_comparison.csv")
    concept = pd.read_csv(PROJECT_ROOT / "data/concept_drift_log.csv")
    overlap = pd.read_csv(PROJECT_ROOT / "data/overlap_results.csv")
    aba_events = pd.read_csv(PROJECT_ROOT / "data/drift_with_aba_events.csv")

    m["comparison"] = comparison
    m["hierarchy"] = hierarchy
    m["concept"] = concept
    m["overlap"] = overlap
    m["aba_events"] = aba_events

    def row(module: str) -> pd.Series:
        return comparison[comparison["module"] == module].iloc[0]

    m["baseline"] = row("Baseline - RouterManager")
    m["disagreement"] = row("Disagreement Module - Borda")
    m["overlap_module"] = row("Overlap Module - LCE+ToM")
    m["drift_module"] = row("Drift Module - EMC+ABA")
    m["interaction"] = row("Interaction Module - Scenario 4")
    m["combined"] = row("Scenario 5 Combined Pass")
    m["post_rl"] = row("Post-RL Reward Loop")
    m["hierarchy_flat"] = hierarchy.iloc[0]
    m["hierarchy_3"] = hierarchy.iloc[1]
    m["concept_stress"] = concept[concept["batch_type"] == "synthetic_stress_test"].iloc[0]
    return m


def comparison_rows(df: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    for _, r in df.iterrows():
        rows.append(
            [
                str(r["module"]),
                str(int(r["profiles_run"])),
                pct(r.get("disagreement_rate")),
                num(r.get("redundancy_index")),
                num(r.get("mean_drift_score")),
                pct(r.get("final_accuracy")),
                str(r["key_finding"]),
            ]
        )
    return rows


def bibliography() -> list[str]:
    return [
        "J.-C. de Borda, 'Memoire sur les elections au scrutin,' Histoire de l'Academie Royale des Sciences, 1781.",
        "D. Dua and C. Graff, UCI Machine Learning Repository, University of California, Irvine, 2019.",
        "UCI Machine Learning Repository, 'Statlog (German Credit Data),' dataset documentation.",
        "S. Kullback and R. A. Leibler, 'On Information and Sufficiency,' Annals of Mathematical Statistics, vol. 22, no. 1, pp. 79-86, 1951.",
        "R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.",
        "G. Salton and C. Buckley, 'Term-weighting approaches in automatic text retrieval,' Information Processing & Management, 1988.",
        "K. Pearson, 'Notes on Regression and Inheritance in the Case of Two Parents,' Proceedings of the Royal Society of London, 1895.",
        "S. Yao et al., 'ReAct: Synergizing Reasoning and Acting in Language Models,' arXiv:2210.03629, 2022.",
        "Q. Wu et al., 'AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation,' arXiv:2308.08155, 2023.",
        "G. Li et al., 'CAMEL: Communicative Agents for Mind Exploration of Large Language Model Society,' arXiv:2303.17760, 2023.",
        "L. Wang et al., 'A Survey on Large Language Model based Autonomous Agents,' arXiv:2308.11432, 2023.",
        "CrewAI, 'CrewAI Documentation,' official documentation for agents, tasks, and crews.",
        "Waseda University / Elsevier Pure, 'Locally Centralized Execution for Less Redundant Computation in Multi-Agent Systems,' publication record.",
        "T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, 2nd ed., Springer, 2009.",
        "H. He and E. Garcia, 'Learning from Imbalanced Data,' IEEE Transactions on Knowledge and Data Engineering, 2009.",
        "Project artefacts: Week 3-5 HDFC multi-agent simulations, frozen CSV outputs, and generated comparison figures.",
    ]


def section5_text(m: dict) -> list[tuple[str, list[str]]]:
    return [
        (
            "5.1 Answering the Research Questions",
            [
                (
                    "RQ1 asked whether structured mechanisms can manage disagreement between specialist credit-risk agents. "
                    f"The results show that disagreement is not a small edge case: the intervention subset produced a final-round disagreement rate of {pct(m['disagreement']['disagreement_rate'])}. "
                    f"Borda Count reached {pct(m['disagreement']['final_accuracy'])} strict accuracy, while CONSENSAGENT was stronger in this run at 55.0%. "
                    "The finding is not that voting alone solves credit risk. The finding is that explicit pre-commitment, discussion capture, and resolution scoring make disagreement measurable instead of hidden inside a final answer."
                ),
                (
                    "RQ2 asked whether overlap can be reduced without removing useful specialist review. The overlap module gave the cleanest operational improvement. "
                    f"The redundancy index fell from 1.1500 to {num(m['overlap_module']['redundancy_index'], 4)}, a 71.0% reduction, while final accuracy on the common subset remained {pct(m['overlap_module']['final_accuracy'])}. "
                    "This supports the value of LCE leader assignment and Theory-of-Mind role prompts in a serial underwriting workflow."
                ),
                (
                    "RQ3 asked whether behavioural drift can be controlled across repeated decisions. EMC reduced checkpoint-40 drift by 11.4%, and EMC plus ABA lowered mean drift further to 0.3555. "
                    "The later hierarchy experiment reduced drift even more, from 0.9782 in the flat two-level baseline to 0.0740 in the three-level governance design. "
                    "This suggests that drift control needs both memory management and escalation architecture."
                ),
            ],
        ),
        (
            "5.2 Governance Gains Versus Accuracy Trade-Offs",
            [
                (
                    "A recurring pattern in the experiments is that governance improvements do not always raise strict Kaggle-label accuracy. "
                    f"The three-level hierarchy reduced disagreement from {pct(m['hierarchy_flat']['disagreement_rate'])} to {pct(m['hierarchy_3']['disagreement_rate'])} and reduced drift by 92.4%, but strict accuracy fell from {pct(m['hierarchy_flat']['final_accuracy'])} to {pct(m['hierarchy_3']['final_accuracy'])}. "
                    "This is not a failure of the hierarchy; it shows that governance layers can become conservative unless thresholds are calibrated."
                ),
                (
                    "The reward loop also produced a sober result. Post-RL accuracy remained at 60.0%, matching the combined pre-RL pipeline but not exceeding it. "
                    "This means the reward loop successfully closed the feedback mechanism, but it could not fully overcome the approval-heavy or refer-heavy tendencies already present in the base agent outputs. "
                    "For a bank, this is a useful result because it prevents overstating the impact of reinforcement-style feedback when the underlying policy quality is still limited."
                ),
            ],
        ),
        (
            "5.3 Why the Causal Chain Matters",
            [
                (
                    "The strongest research contribution is the causal-chain framing: Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk. "
                    "The Week 4 interaction run showed 100.0% co-occurrence between overlap and disagreement on the common subset. "
                    "It also showed that resolution can induce additional invocations, which means a disagreement solver can create new overlap unless task ownership is controlled."
                ),
                (
                    "This has direct implications for HDFC-style workflows. A loan pipeline can look orderly at the final-decision level while still containing hidden repeated checks, unstable specialist behaviour, and inconsistent escalation. "
                    "Therefore the system should be evaluated using process metrics such as disagreement rate, redundancy index, drift score, sycophancy flags, and hierarchy reset events, not only final approval accuracy."
                ),
            ],
        ),
        (
            "5.4 Limitations",
            [
                (
                    "The project uses the German Credit dataset and simulated local LLM agents rather than HDFC internal borrower files. This protects confidentiality but limits direct production claims. "
                    "The dataset also does not contain real supporting documents such as bank statements, KYC files, employer letters, or salary slips, so document-level verification is approximated through structured fields."
                ),
                (
                    "The intervention modules were often evaluated on 20-profile or 40-profile subsets because local LLM execution was slow. "
                    "The baseline and hierarchy layers use 50 profiles, and concept-drift monitoring uses 200 natural profiles plus a synthetic stress test. "
                    "Future runs should scale the integrated pipeline to 100 or 200 profiles once runtime constraints are reduced."
                ),
            ],
        ),
    ]


def section6_text() -> list[tuple[str, list[str]]]:
    return [
        (
            "6.1 Conclusion",
            [
                (
                    "This project built and evaluated a multi-agent credit-risk simulation for HDFC-style loan review. "
                    "The work shows that multi-agent systems should not be judged only by whether the final answer is correct. "
                    "They also need to be measured for disagreement, redundant work, behavioural drift, concept drift, and governance traceability."
                ),
                (
                    "Across the experiments, the most reliable improvement came from process control. LCE and ToM reduced redundant work by 71.0%. "
                    "EMC and ABA reduced behavioural drift, and the three-level hierarchy reduced checkpoint drift by 92.4% compared with the flat two-level baseline. "
                    "The disagreement module made private pre-commit positions observable, and the sycophancy detector highlighted cases where agents changed position while becoming too similar to one another."
                ),
                (
                    "The accuracy results are intentionally reported conservatively. Some modules improved strict accuracy relative to the frozen common-subset baseline, while others improved governance at the cost of strict approve/reject accuracy. "
                    "This is an important practical lesson: in regulated lending, a safer and more traceable decision pipeline may sometimes produce more referrals until thresholds are tuned."
                ),
            ],
        ),
        (
            "6.2 Future Work",
            [
                (
                    "The immediate next step is to run the integrated pipeline on a larger 100-200 profile sample and replace local simulated evidence with richer document-like inputs. "
                    "A stronger version would include a calibrated final decision layer that maps Approve, Refer, and Reject into bank-specific risk policy thresholds instead of treating all non-matches as simple errors."
                ),
                (
                    "A second extension is to convert the current reward loop into a safer policy-learning layer. Instead of only changing module weights after Kaggle feedback, the system should learn when to escalate, when to request review, and when to preserve independent specialist disagreement. "
                    "This would keep the governance benefits while reducing approval bias and excessive referral."
                ),
            ],
        ),
    ]


def write_section_docx(title: str, filename: str, sections: list[tuple[str, list[str]]], summary: str) -> None:
    doc = Document()
    style_document(doc, title)
    add_title(
        doc,
        title,
        "HDFC multi-agent credit-risk simulation",
        [
            ("Prepared for", "Week 5 Thursday research package"),
            ("Prepared by", "Anil Kadam"),
            ("Date", date.today().isoformat()),
            ("Status", "Draft for report integration"),
        ],
    )
    doc.add_page_break()
    doc.add_heading(title, level=1)
    add_callout(doc, "Section purpose", summary)
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            add_para(doc, para)
    doc.add_heading("References", level=1)
    for item in bibliography():
        add_para(doc, item)
    path = OUT_DIR / filename
    doc.save(path)

    md = ["# " + title, "", summary, ""]
    for heading, paragraphs in sections:
        md.extend(["## " + heading, ""])
        for para in paragraphs:
            md.extend([para, ""])
    md.extend(["## References", ""])
    md.extend([f"- {item}" for item in bibliography()])
    (OUT_DIR / filename.replace(".docx", ".md")).write_text("\n".join(md), encoding="utf-8")


def write_summary_docx(m: dict) -> None:
    doc = Document()
    style_document(doc, "HDFC 5-Scenario Summary")
    add_title(
        doc,
        "HDFC 5-Scenario Summary",
        "Manager-ready summary of the multi-agent credit-risk simulation",
        [
            ("Coverage", "Disagreement, overlap, drift, interaction, integrated governance"),
            ("Prepared by", "Anil Kadam"),
            ("Date", date.today().isoformat()),
            ("Evidence base", "Frozen CSV outputs and generated figures"),
        ],
    )
    doc.add_page_break()
    doc.add_heading("Executive Snapshot", level=1)
    add_callout(
        doc,
        "Bottom line",
        "The project does not claim that every intervention improves strict accuracy. It shows which governance controls reduce hidden multi-agent failure modes such as disagreement, redundant checks, and drift.",
    )

    rows = [
        [
            "Scenario 1 - Disagreement",
            "Agents gave conflicting decisions on the same borrower profile.",
            "Pre-commit, one discussion round, Borda Count, CONSENSAGENT, cosine sycophancy detector.",
            "Final disagreement 100.0%; Borda accuracy 40.0%; CONSENSAGENT accuracy 55.0%; 13 sycophancy flags.",
        ],
        [
            "Scenario 2 - Overlap",
            "Agents repeated the same repayment-capacity and evidence checks.",
            "LCE leader assignment, TAP intent board, redundancy scorer, ToM role prompts.",
            "Redundancy index reduced from 1.1500 to 0.3333; duplicate invocations reduced by 71.0%.",
        ],
        [
            "Scenario 3 - Drift",
            "Repeated sequential reviews pushed agent distributions away from the frozen baseline.",
            "EMC memory compaction, KL drift scoring, ABA re-anchoring.",
            "EMC reduced drift by 11.4%; EMC+ABA reduced checkpoint drift to 0.3555.",
        ],
        [
            "Scenario 4 - Interaction Chain",
            "Solving one failure mode can trigger another.",
            "Causal chain test: Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk.",
            "Overlap and disagreement co-occurred in 100.0% of the common subset; LCE removed 2.45 excess invocations per profile.",
        ],
        [
            "Scenario 5 - Integrated Governance",
            "A combined pipeline needs both feedback and escalation control.",
            "Combined pass, RL reward loop, concept drift layer, 3-level L1/L2/L3 hierarchy.",
            "Post-RL accuracy 60.0%; concept drift stress test detected r=0.3905; hierarchy reduced drift by 92.4%.",
        ],
    ]
    add_table(doc, ["Scenario", "Problem", "Mechanism", "Observed Outcome"], rows, widths=[1.45, 1.55, 2.05, 2.25])

    doc.add_heading("How to Read the Results", level=1)
    add_bullets(
        doc,
        [
            "Accuracy is one outcome, but not the only useful outcome in a regulated credit-risk workflow.",
            "Redundancy index measures repeated work; lower is better when specialist coverage is preserved.",
            "Drift score measures how far an agent's decision distribution moved from the frozen baseline.",
            "Refer outcomes should be treated as governance escalations, not always as simple model errors.",
            "The current numbers are simulation results, not production lending performance claims.",
        ],
    )

    path = OUT_DIR / "HDFC_5Scenario_Summary.docx"
    doc.save(path)

    md_lines = ["# HDFC 5-Scenario Summary", "", "Manager-ready summary of the simulation.", ""]
    md_lines.append("| Scenario | Problem | Mechanism | Observed Outcome |")
    md_lines.append("|---|---|---|---|")
    for row in rows:
        md_lines.append("| " + " | ".join(row) + " |")
    (OUT_DIR / "HDFC_5Scenario_Summary.md").write_text("\n".join(md_lines), encoding="utf-8")


def write_full_report_docx(m: dict) -> None:
    intro = strip_md_heading(read_text(HDFC_ROOT / "Week5_Monday/Section_1_Introduction.md"))
    cross = strip_md_heading(read_text(HDFC_ROOT / "Week5_Monday/Section_2_4_CrossPhenomenon.md"))
    concept = strip_md_heading(read_text(HDFC_ROOT / "Week5_Monday/Section_3_ConceptDrift_Addition.md"))
    results = strip_md_heading(read_text(HDFC_ROOT / "Week5_Tuesday/Section_4_Results_Complete.md"))
    hierarchy = strip_md_heading(read_text(HDFC_ROOT / "Week5_Wednesday/Section_3Level_Hierarchy.md"))

    doc = Document()
    style_document(doc, "Research Report Draft v1")
    add_title(
        doc,
        "Multi-Agent Governance for Credit-Risk Simulation",
        "Disagreement, overlap, drift, reward feedback, concept drift, and hierarchy in an HDFC-style loan review pipeline",
        [
            ("Author", "Anil Kadam"),
            ("Draft", "Week 5 Thursday - Research Report Draft v1"),
            ("Dataset", "German Credit research dataset"),
            ("Status", "Integrated draft for mentor review"),
        ],
    )
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    for para in [
        "This report presents a simulated multi-agent credit-risk pipeline inspired by HDFC-style underwriting operations. Seven agents represent routing, income, fraud, credit, compliance, summarisation, and weak-model comparison roles.",
        "The project studies five scenarios: disagreement management, overlap control, agent drift control, cross-phenomenon interaction, and integrated governance with reward feedback and hierarchy.",
        "The central finding is that multi-agent credit systems require process governance, not only final-decision accuracy. The strongest improvements were reductions in redundancy and drift, while strict accuracy remained sensitive to approval bias, referral policy, and threshold calibration.",
    ]:
        add_para(doc, para)

    doc.add_heading("Key Metrics", level=2)
    add_table(
        doc,
        ["Module", "Profiles", "Disagreement", "Redundancy", "Drift", "Accuracy", "Finding"],
        comparison_rows(m["comparison"]),
        widths=[1.35, 0.55, 0.75, 0.75, 0.65, 0.65, 2.55],
    )

    doc.add_page_break()
    doc.add_heading("1. Introduction", level=1)
    for para in paragraph_chunks(intro):
        add_para(doc, para)

    doc.add_heading("2. Background and Literature", level=1)
    doc.add_heading("2.1 Multi-Agent Credit-Risk Workflows", level=2)
    for para in [
        "Credit-risk decisions are naturally multi-stage. A borrower can be assessed through income stability, fraud signals, credit behaviour, compliance completeness, and final routing. Multi-agent systems map well to this structure because each agent can specialise in one reasoning domain.",
        "However, specialisation introduces coordination risk. If every agent independently reviews the same evidence, the system becomes redundant. If agents observe one another too early, disagreement may collapse into imitation. If agents carry too much history, their decisions may drift away from the original baseline.",
    ]:
        add_para(doc, para)
    doc.add_heading("2.2 Voting, Consensus, and Sycophancy", level=2)
    for para in [
        "Borda Count is used as a structured voting mechanism because it uses ranked preferences rather than only plurality winners. In this project, rank 1 receives 2 points, rank 2 receives 1 point, and rank 3 receives 0 points across Approve, Refer, and Reject.",
        "CONSENSAGENT is used as a complementary confidence-weighted mechanism. It rewards groups that are both confident and stable from their pre-commit positions. The cosine sycophancy detector then checks whether an agent changed position while becoming too textually similar to another agent's reasoning.",
    ]:
        add_para(doc, para)
    doc.add_heading("2.3 Overlap, Drift, and Memory", level=2)
    for para in [
        "The overlap module uses locally centralised execution (LCE), task-intent announcements, and Theory-of-Mind prompting to prevent agents from repeating the same check. Drift is monitored using KL-style distribution shift between current agent decisions and frozen baseline decisions.",
        "Episodic Memory Consolidation compresses long agent memory into short summaries. Adaptive Behaviour Anchoring reintroduces baseline examples when drift exceeds a threshold. Together, these mechanisms separate context control from active behavioural correction.",
    ]:
        add_para(doc, para)
    doc.add_heading("2.4 Cross-Phenomenon Interaction", level=2)
    for para in paragraph_chunks(cross):
        add_para(doc, para)

    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Dataset and Agent Roles", level=2)
    add_para(
        doc,
        "The simulation uses the German Credit dataset as a public, non-confidential proxy for credit-risk review. Borrower profiles are built from structured fields such as age, job type, housing, savings, checking status, credit amount, duration, and loan purpose.",
    )
    add_table(
        doc,
        ["Agent", "Primary Role"],
        [
            ["RouterManager", "Coordinates final routing and overall credit-risk path."],
            ["IncomeAgent", "Assesses repayment capacity using available financial proxy fields."],
            ["FraudAgent", "Checks suspicious patterns and missing-evidence risk."],
            ["CreditAgent", "Reviews credit behaviour and repayment burden."],
            ["ComplianceAgent", "Checks policy fit and KYC-style completeness."],
            ["SummariserAgent", "Compresses multi-agent evidence into decision-ready summaries."],
            ["WeakModelAgent", "Provides a low-capability comparison view to expose blind spots."],
        ],
        widths=[1.6, 5.4],
    )
    doc.add_heading("3.2 Scenario Modules", level=2)
    add_bullets(
        doc,
        [
            "Scenario 1: Pre-commit disagreement, discussion, Borda Count, CONSENSAGENT, and sycophancy detection.",
            "Scenario 2: LCE leader assignment, TAP intent board, redundancy scoring, TRACE bidding, and ToM prompts.",
            "Scenario 3: EMC memory compaction, KL-style drift scoring, and ABA re-anchoring.",
            "Scenario 4: Causal chain between overlap, disagreement, resolution, renewed overlap, and drift risk.",
            "Scenario 5: Combined governance pass with reward-loop feedback and later hierarchy control.",
        ],
    )
    doc.add_heading("3.3 Concept Drift Correlation Layer", level=2)
    for para in paragraph_chunks(concept):
        add_para(doc, para)
    doc.add_heading("3.4 Three-Level Governance Hierarchy", level=2)
    for para in paragraph_chunks(hierarchy):
        add_para(doc, para)

    doc.add_heading("4. Results", level=1)
    for para in paragraph_chunks(results):
        if para.startswith("Section 4.1"):
            doc.add_heading("4.1 Disagreement", level=2)
            continue
        if para.startswith("Section 4.2"):
            doc.add_heading("4.2 Overlap", level=2)
            continue
        if para.startswith("Section 4.3"):
            doc.add_heading("4.3 Drift", level=2)
            continue
        if para.startswith("Section 4.4"):
            doc.add_heading("4.4 Causal Chain", level=2)
            continue
        if para.startswith("Section 4.5"):
            doc.add_heading("4.5 Combined Pipeline and Reward Loop", level=2)
            continue
        add_para(doc, para)

    doc.add_heading("4.6 Visual Evidence", level=2)
    add_figure(doc, PROJECT_ROOT / "results/Figure_1_Disagreement_Rate.png", "Figure 1. Disagreement and resolution outcomes.")
    add_figure(doc, PROJECT_ROOT / "results/Figure_2_Redundancy_Reduction.png", "Figure 2. Redundancy reduction after LCE and ToM.")
    add_figure(doc, PROJECT_ROOT / "results/Figure_4A_ABA_Drift_Comparison.png", "Figure 3. Drift reduction across no-EMC, EMC, and EMC+ABA conditions.")
    add_figure(doc, PROJECT_ROOT / "results/Figure_6_Concept_Drift_Correlation.png", "Figure 4. Concept drift correlation layer with synthetic stress test.")
    add_figure(doc, PROJECT_ROOT / "results/Figure_10_Hierarchy_Comparison.png", "Figure 5. Two-level versus three-level hierarchy comparison.")

    doc.add_heading("5. Discussion", level=1)
    for heading, paragraphs in section5_text(m):
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            add_para(doc, para)

    doc.add_heading("6. Conclusion", level=1)
    for heading, paragraphs in section6_text():
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            add_para(doc, para)

    doc.add_heading("Bibliography", level=1)
    for item in bibliography():
        add_para(doc, item)

    path = OUT_DIR / "Research_Report_Draft_v1.docx"
    doc.save(path)

    md = [
        "# Multi-Agent Governance for Credit-Risk Simulation",
        "",
        "Integrated Week 5 Thursday report draft. See the DOCX for tables and embedded figures.",
        "",
        "## Key Metrics",
        "",
        m["comparison"].to_markdown(index=False),
        "",
        "## Bibliography",
        "",
    ]
    md.extend([f"- {item}" for item in bibliography()])
    (OUT_DIR / "Research_Report_Draft_v1.md").write_text("\n".join(md), encoding="utf-8")


def write_bibliography_md() -> None:
    lines = ["# Bibliography", ""]
    for i, item in enumerate(bibliography(), 1):
        lines.append(f"{i}. {item}")
    (OUT_DIR / "Bibliography.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_summary_md(m: dict) -> None:
    lines = [
        "# Week 5 Thursday Package Summary",
        "",
        "Generated outputs:",
        "",
        "- Section_5_Discussion.docx",
        "- Section_6_Conclusion.docx",
        "- HDFC_5Scenario_Summary.docx",
        "- Research_Report_Draft_v1.docx",
        "",
        "Headline metrics:",
        "",
        f"- Disagreement module: Borda accuracy {pct(m['disagreement']['final_accuracy'])}, final disagreement {pct(m['disagreement']['disagreement_rate'])}.",
        f"- Overlap module: redundancy index {num(m['overlap_module']['redundancy_index'], 4)}, reduction 71.0%.",
        f"- Drift module: EMC+ABA mean drift {num(m['drift_module']['mean_drift_score'], 4)}, final accuracy {pct(m['drift_module']['final_accuracy'])}.",
        f"- Post-RL reward loop: final accuracy {pct(m['post_rl']['final_accuracy'])}.",
        f"- Hierarchy: drift reduced by 92.4%, but strict accuracy fell to {pct(m['hierarchy_3']['final_accuracy'])}.",
        f"- Concept drift stress test: correlation r={num(m['concept_stress']['correlation_r'], 4)}, drift detected.",
        "",
        "Interpretation:",
        "",
        "The project shows process-governance gains more strongly than pure accuracy gains. This is a useful and honest research result for regulated lending.",
    ]
    (OUT_DIR / "Week5_Thursday_Summary.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    m = load_metrics()
    write_section_docx(
        "Section 5 - Discussion",
        "Section_5_Discussion.docx",
        section5_text(m),
        "This section interprets the results across all modules, connects them to the research questions, and states limitations honestly.",
    )
    write_section_docx(
        "Section 6 - Conclusion",
        "Section_6_Conclusion.docx",
        section6_text(),
        "This section gives the final takeaway, contribution statement, and future-work path for the HDFC multi-agent credit-risk project.",
    )
    write_summary_docx(m)
    write_full_report_docx(m)
    write_bibliography_md()
    write_summary_md(m)
    print(f"Generated Week 5 Thursday package in: {OUT_DIR}")
    for p in sorted(OUT_DIR.glob("*")):
        print(f"- {p.name}")


if __name__ == "__main__":
    main()
