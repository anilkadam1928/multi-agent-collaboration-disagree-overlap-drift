from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = PROJECT_ROOT / "analysis"
RESULTS_DIR = PROJECT_ROOT / "results"
DATA_DIR = PROJECT_ROOT / "data"

MASTER_TABLE_PATH = ANALYSIS_DIR / "master_comparison_table_week4.csv"

FINAL_REPORT_DOCX = ANALYSIS_DIR / "Week4_Friday_Final_Package.docx"
SECTION3_MD = ANALYSIS_DIR / "Week4_Friday_Section3_Draft.md"
SECTION3_DOCX = ANALYSIS_DIR / "Week4_Friday_Section3_Draft.docx"
MENTOR_MD = ANALYSIS_DIR / "Week4_Friday_Mentor_Checkin_Summary.md"
GITHUB_MD = ANALYSIS_DIR / "Week4_Friday_GitHub_Checklist.md"
SUBMISSION_INDEX_MD = ANALYSIS_DIR / "Week4_Friday_Submission_Index.md"

FIGURES = [
    ("Figure 1", "Disagreement Rate", RESULTS_DIR / "Figure_1_Disagreement_Rate.png"),
    ("Figure 2", "Redundancy Reduction", RESULTS_DIR / "Figure_2_Redundancy_Reduction.png"),
    ("Figure 3", "Drift Score With EMC and ABA", RESULTS_DIR / "Figure_3_Drift_Score_With_ABA.png"),
    ("Figure 4", "Overlap-Disagreement Causal Chain", RESULTS_DIR / "Figure_4_Causal_Chain.png"),
    ("Figure 5", "Master Comparison", RESULTS_DIR / "Figure_5_Master_Comparison.png"),
]


def pct(value) -> str:
    if value == "" or pd.isna(value):
        return ""
    return f"{float(value) * 100:.1f}%"


def num(value, digits: int = 3) -> str:
    if value == "" or pd.isna(value):
        return ""
    return f"{float(value):.{digits}f}"


def load_master() -> pd.DataFrame:
    if not MASTER_TABLE_PATH.exists():
        raise FileNotFoundError(f"Missing master table: {MASTER_TABLE_PATH}")
    return pd.read_csv(MASTER_TABLE_PATH)


def get_metric(master: pd.DataFrame, module: str, column: str):
    row = master[master["module"] == module]
    if row.empty:
        return ""
    return row.iloc[0][column]


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_small_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(95, 95, 95)


def add_master_table(doc: Document, master: pd.DataFrame) -> None:
    doc.add_heading("Master Comparison Table", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Module", "Profiles", "Disagreement", "Redundancy", "Drift", "Accuracy"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for _, row in master.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["module"])
        cells[1].text = str(int(row["profiles_run"]))
        cells[2].text = pct(row["disagreement_rate"])
        cells[3].text = num(row["redundancy_index"])
        cells[4].text = num(row["mean_drift_score"])
        cells[5].text = pct(row["final_accuracy"])

    doc.add_paragraph()
    add_small_note(
        doc,
        "Note: Empty cells mean the metric is not the primary measure for that module. "
        "All intervention comparisons use the common B001-B020 subset unless otherwise stated.",
    )


def add_figures(doc: Document) -> None:
    doc.add_heading("Figures For Report Inclusion", level=1)
    for fig_id, title, path in FIGURES:
        if not path.exists():
            continue
        doc.add_heading(f"{fig_id}: {title}", level=2)
        doc.add_picture(str(path), width=Inches(6.2))
        caption = doc.add_paragraph(f"{fig_id}. {title}. Source: {path.name}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)


def write_section3_markdown(master: pd.DataFrame) -> None:
    baseline_acc = pct(get_metric(master, "Baseline - RouterManager", "final_accuracy"))
    baseline_dis = pct(get_metric(master, "Baseline - RouterManager", "disagreement_rate"))
    borda_acc = pct(get_metric(master, "Disagreement Module - Borda", "final_accuracy"))
    overlap_acc = pct(get_metric(master, "Overlap Module - LCE+ToM", "final_accuracy"))
    redundancy = num(get_metric(master, "Overlap Module - LCE+ToM", "redundancy_index"))
    drift = num(get_metric(master, "Drift Module - EMC+ABA", "mean_drift_score"))
    drift_acc = pct(get_metric(master, "Drift Module - EMC+ABA", "final_accuracy"))
    combined_acc = pct(get_metric(master, "Scenario 5 Combined Pass", "final_accuracy"))

    text = f"""# Section 3 Draft - Week 4 Results and Analysis

## 3.1 Experimental Setup

The Week 4 evaluation used the frozen Week 3 baseline as the reference point. The baseline file contains 50 German Credit profiles, while the intervention modules were evaluated on the common B001-B020 subset because local LLM runtime constraints made repeated 50-profile intervention runs expensive. All cross-module comparisons in this section are therefore computed on the same borrower IDs, which keeps the comparison internally consistent.

The pipeline evaluates five intervention layers: disagreement resolution, overlap control, drift reduction, interaction analysis, and a combined Scenario 5 integration pass. The objective is not only to maximize final accuracy, but also to measure whether the multi-agent system becomes more stable, less redundant, and less prone to drift.

## 3.2 Baseline Behaviour

The frozen RouterManager baseline achieved {baseline_acc} accuracy on the common subset. The baseline also showed a high agent disagreement rate of {baseline_dis}. This confirms the initial research problem: an unconstrained multi-agent lending pipeline can produce inconsistent specialist outputs and unstable final routing.

## 3.3 Disagreement Module

The disagreement module introduced private pre-commit decisions, a discussion round, Borda Count voting, CONSENSAGENT scoring, and cosine-based sycophancy detection. The key methodological point is that each agent first states its independent position before seeing other agents' outputs. This creates a baseline position for detecting later shifts.

Using Borda Count as the primary resolution mechanism, the module achieved {borda_acc} accuracy. The result improved over the baseline but did not fully solve decision quality. The main contribution is therefore methodological: the module exposes disagreement, measures position changes, and prevents group discussion from being treated as a black box.

## 3.4 Overlap Module

The overlap module implemented LCE leader assignment, TAP intent boards, redundancy scoring, TRACE priority bidding, and Theory-of-Mind role prompts. The system models asynchronous serial overlap, meaning agents act stage-by-stage rather than in simultaneous parallel execution.

The controlled LCE+ToM configuration reached a redundancy index of {redundancy} and final accuracy of {overlap_acc}. Compared with the raw overlap setting, redundancy fell by 71.0%. This is one of the strongest Week 4 findings: leadership and role clarity substantially reduce duplicate or overlapping checks.

## 3.5 Drift Module

The drift module measured KL-divergence between current agent decision distributions and the frozen baseline. EMC compresses long agent memory into short summaries, while ABA injects anchor examples when drift exceeds the threshold.

The final EMC+ABA drift score was {drift}, with final accuracy of {drift_acc}. Compared with EMC-only drift, ABA reduced drift by 59.0%. This supports the hypothesis that memory compaction plus anchoring can stabilize long-running agent behaviour.

## 3.6 Interaction Module

The interaction analysis linked overlap and disagreement as a causal chain: overlap can create disagreement, disagreement resolution can create additional invocations, and those repeated invocations can amplify drift risk. On the common subset, overlap and disagreement co-occurred in 100.0% of cases.

Raw resolution produced an average of 3.45 duplicate or recheck invocations. After LCE+ToM, remaining invocations fell to 1.00. This shows that overlap control is not only an efficiency mechanism; it also reduces downstream instability after disagreement resolution.

## 3.7 Combined Scenario 5 Pass

The combined Scenario 5 integration pass merged the available outputs from the baseline, disagreement, overlap, drift, and interaction modules. The integrated pass achieved {combined_acc} final accuracy on the common subset.

This result should be interpreted carefully. The combined decision layer still inherits approval bias from the RouterManager-led route in several cases. However, the Week 4 contribution is clear: the system now logs disagreement, controls redundancy, measures drift, applies anchoring, and produces a reward-ready decision trace for Week 5 reinforcement learning.

## 3.8 Limitations

The primary limitation is sample size. Intervention modules were evaluated on 20 common profiles, while the frozen baseline contains 50 profiles. This was due to local LLM runtime constraints. The next stage should run the integrated pipeline on 100 or 200 profiles once prompts and modules are frozen.

The second limitation is decision bias. Some agents still over-approve or produce unknown outputs. This should be treated as a finding rather than hidden: the architecture improves stability and traceability, but final credit quality still needs calibration.

## 3.9 Summary

Week 4 successfully moved the project from isolated modules to a measurable multi-agent control framework. The strongest quantitative results are the 71.0% redundancy reduction from LCE+ToM and the 59.0% drift reduction from ABA over EMC-only. The combined pipeline reached {combined_acc} accuracy and produced the reward logs needed for Week 5.
"""
    SECTION3_MD.write_text(text, encoding="utf-8")


def markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    doc = Document()
    text = markdown_path.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    doc.save(docx_path)


def write_mentor_summary(master: pd.DataFrame) -> None:
    text = f"""# Week 4 Mentor Check-in Summary

## What is complete

- Disagreement module: pre-commit, discussion, Borda Count, CONSENSAGENT, and sycophancy detection.
- Overlap module: LCE leader assignment, TAP intent board, redundancy scorer, TRACE bidding, and ToM prompts.
- Drift module: EMC memory compaction, KL drift scoring, and ABA drift correction.
- Interaction module: causal chain from overlap to disagreement to renewed overlap and drift risk.
- Combined Scenario 5 pass: integrated output table, master comparison chart, and RL reward stub.

## Key numbers

- Baseline RouterManager accuracy: 25.0% on the common 20-profile subset.
- Disagreement/Borda accuracy: 40.0%.
- LCE+ToM final accuracy: 60.0%.
- Redundancy reduction with LCE+ToM: 71.0%.
- EMC-only drift reduction: 11.4%.
- ABA reduction vs EMC-only: 59.0%.
- Combined Scenario 5 accuracy: 60.0%.

## Honest limitation

The intervention modules were evaluated on B001-B020 due to local LLM runtime constraints. The baseline remains frozen at 50 profiles. Larger 100/200-profile runs are planned for Week 5 after the modules are stable.

## Main interpretation

The strongest result is not raw accuracy alone. The project now has measurable controls for disagreement, overlap, drift, and reward logging. The system is more inspectable and less redundant, but final decision calibration still needs improvement because some routes remain approval-biased.
"""
    MENTOR_MD.write_text(text, encoding="utf-8")


def write_github_checklist() -> None:
    text = """# Week 4 GitHub Push Checklist

## Files to include

- agents.py
- config.py
- data_loader.py
- profile_builder.py
- output_parser.py
- results_logger.py
- disagreement_simulation.py
- overlap_simulation.py
- drift_simulation.py
- drift_aba_simulation.py
- interaction_simulation.py
- combined_pipeline.py
- anchor_examples.py

## Results to include

- data/baseline_results.csv
- results/disagreement_results.csv
- data/overlap_results.csv
- data/drift_with_emc.csv
- data/drift_with_aba.csv
- data/drift_with_aba_events.csv
- data/dual_results.csv
- data/combined_results.csv
- data/rl_rewards.csv
- analysis/master_comparison_table_week4.csv

## Figures to include

- results/Figure_1_Disagreement_Rate.png
- results/Figure_2_Redundancy_Reduction.png
- results/Figure_3_Drift_Score_With_ABA.png
- results/Figure_4_Causal_Chain.png
- results/Figure_5_Master_Comparison.png

## Reports to include

- analysis/Week4_Day1_Drift_Report.docx
- analysis/Week4_Day2_ABA_Report.docx
- analysis/Week4_Day3_Interaction_Report.docx
- analysis/Week4_Day4_Combined_Report.docx
- analysis/Week4_Friday_Final_Package.docx
- analysis/Week4_Friday_Section3_Draft.docx

## Suggested commit message

Week 4 multi-agent controls: disagreement, overlap, drift, ABA, and combined analysis

## Note

This folder is currently not detected as a Git repository on this machine. If you want to push to GitHub, initialize or open the correct repository folder first.
"""
    GITHUB_MD.write_text(text, encoding="utf-8")


def write_submission_index() -> None:
    lines = [
        "# Week 4 Friday Submission Index",
        "",
        "## Final Friday files",
        f"- {FINAL_REPORT_DOCX}",
        f"- {SECTION3_DOCX}",
        f"- {SECTION3_MD}",
        f"- {MENTOR_MD}",
        f"- {GITHUB_MD}",
        "",
        "## Core Week 4 reports",
        f"- {ANALYSIS_DIR / 'Week4_Day1_Drift_Report.docx'}",
        f"- {ANALYSIS_DIR / 'Week4_Day2_ABA_Report.docx'}",
        f"- {ANALYSIS_DIR / 'Week4_Day3_Interaction_Report.docx'}",
        f"- {ANALYSIS_DIR / 'Week4_Day4_Combined_Report.docx'}",
        "",
        "## Core figures",
    ]
    for _, title, path in FIGURES:
        lines.append(f"- {title}: {path}")
    lines.append("")
    lines.append("## Master comparison")
    lines.append(f"- {MASTER_TABLE_PATH}")
    SUBMISSION_INDEX_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_final_report(master: pd.DataFrame) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(
        doc,
        "Week 4 Friday Final Package",
        "Output analysis, comparison table, Section 3 draft, and mentor check-in summary",
    )
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    summary_points = [
        "Week 4 completed the control layer around the HDFC multi-agent lending pipeline.",
        "The strongest quantitative findings are 71.0% redundancy reduction from LCE+ToM and 59.0% drift reduction from ABA vs EMC-only.",
        "The combined Scenario 5 pass achieved 60.0% final accuracy on the common 20-profile subset.",
        "The main remaining limitation is decision calibration: some final routes remain approval-biased.",
    ]
    for point in summary_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Scope Statement", level=1)
    doc.add_paragraph(
        "Intervention modules were evaluated on a common 20-profile subset drawn from the frozen 50-profile baseline. "
        "This is acceptable for Week 4 because all intervention comparisons use the same borrower IDs. "
        "Larger 100/200-profile runs should be deferred to Week 5 after prompts and modules are frozen."
    )

    add_master_table(doc, master)

    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "The Week 4 results show that the architecture improves inspectability and control even when final accuracy is not yet production-grade. "
        "Disagreement is now explicit, overlap is measurable, drift is quantified, and the combined pass creates a reward-ready log for Week 5."
    )
    doc.add_paragraph(
        "The approval bias in the final combined decision layer should be reported honestly. It is not a failure of the week; it is the next calibration target."
    )

    add_figures(doc)

    doc.add_heading("Friday Completion Checklist", level=1)
    checklist = [
        "Master comparison table completed.",
        "Section 3 draft prepared.",
        "Mentor check-in summary prepared.",
        "GitHub checklist prepared.",
        "Final package document prepared.",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(FINAL_REPORT_DOCX)


def main() -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    master = load_master()
    write_section3_markdown(master)
    markdown_to_docx(SECTION3_MD, SECTION3_DOCX)
    write_mentor_summary(master)
    write_github_checklist()
    write_submission_index()
    write_final_report(master)

    print("Friday package generated.")
    for path in [
        FINAL_REPORT_DOCX,
        SECTION3_DOCX,
        SECTION3_MD,
        MENTOR_MD,
        GITHUB_MD,
        SUBMISSION_INDEX_MD,
    ]:
        print(path)


if __name__ == "__main__":
    main()
