from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "generated_report_sections" / "Week5_Monday"
SECTION_1_DOCX = OUTPUT_DIR / "Section_1_Introduction.docx"
SECTION_1_MD = OUTPUT_DIR / "Section_1_Introduction.md"
SECTION_24_DOCX = OUTPUT_DIR / "Section_2_4_CrossPhenomenon.docx"
SECTION_24_MD = OUTPUT_DIR / "Section_2_4_CrossPhenomenon.md"


SECTION_1_TITLE = "Section 1: Introduction"
SECTION_24_TITLE = "Section 2.4: Cross-Phenomenon Interaction"


SECTION_1_TEXT = """Credit risk decision-making is one of the most consequential areas where financial institutions are beginning to evaluate multi-agent AI systems. In a bank such as HDFC Bank, a loan application is not judged through a single signal. It passes through a sequence of specialist checks: income stability, repayment capacity, fraud indicators, credit behaviour, compliance completeness, and final approval routing. A multi-agent architecture is a natural fit for this workflow because each agent can represent a specialist function. However, the same structure also creates new failure modes that are not visible in a single-model system.

This project studies three such failure modes in a simulated credit-risk pipeline. The first is agent disagreement: different agents may interpret the same borrower profile differently and reach conflicting decisions such as approve, reject, or refer. The second is agent overlap: multiple agents may redundantly perform the same check, increasing cost and creating repeated evidence loops. The third is agent drift: over many sequential loan reviews, agents may gradually move away from their original decision standard and become too lenient, too strict, or inconsistent.

The research is framed around three questions. RQ1 asks: when and why do agents disagree during credit-risk evaluation, and can structured resolution mechanisms such as Borda Count improve decision consistency? RQ2 asks: when do agents overlap in a serial loan-review pipeline, and can leader assignment, task-intent announcements, and role clarity reduce redundant work? RQ3 asks: how does agent behaviour drift over repeated decisions, and can memory compaction and adaptive anchoring reduce that drift?

The project uses the public German Credit dataset rather than real bank data, preserving confidentiality while still modelling a realistic underwriting-style workflow. The simulation is implemented in Python using a local multi-agent setup, with specialist agents representing routing, income, fraud, credit, compliance, summarisation, and weak-model comparison roles. Results are compared against a frozen baseline so that changes in disagreement, redundancy, drift, and final accuracy can be measured consistently.

The novel contribution of this project is the causal chain connecting the three phenomena: Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk. Instead of treating disagreement, overlap, and drift as isolated problems, the simulation studies how they interact inside the same credit-risk pipeline. This matters for HDFC-style risk operations because an intervention that resolves one failure mode may accidentally amplify another unless the full chain is monitored."""


SECTION_24_TEXT = """Existing multi-agent research often studies disagreement, overlap, or drift as separate behaviours. Disagreement is usually treated as a voting or consensus problem. Overlap is usually treated as redundant task execution. Drift is usually treated as a long-run calibration problem. In an actual credit-risk workflow, however, these behaviours can interact. A loan application moves through multiple stages, and decisions made at one stage influence which agents act again later. This creates a causal chain rather than three isolated events.

The interaction model used in this project is: Overlap -> Disagreement -> Resolution -> Renewed Overlap -> Drift Risk. The first link occurs when two or more agents inspect similar evidence or duplicate the same check. Because each agent has a different role and prompt framing, redundant inspection does not guarantee identical interpretation. Instead, overlap can expose conflicting judgements. For example, one agent may view low savings as a credit concern, while another may treat stable housing as enough to approve. This converts overlap into disagreement.

The second link is disagreement -> resolution. Once agents disagree, the system must resolve the conflict through a mechanism such as Borda Count or CONSENSAGENT. Resolution is useful, but it is not free. In a serial pipeline, resolution can trigger additional review steps, re-checks, or summary passes. This creates the third link: resolution -> renewed overlap. The system may ask agents to reconsider, explain, compare, or validate their decisions, causing new overlapping work even after the original conflict has been resolved.

The final link is renewed overlap -> drift risk. Repeated reconsideration can change an agent's working context and decision distribution. If agents repeatedly see group decisions, summary outputs, or previous justifications, they may gradually move away from their original baseline behaviour. This is why the project connects the interaction module to the drift module. The same mechanism that resolves disagreement can also add repeated context, and repeated context can contribute to drift.

No existing paper models this causal chain. This simulation provides the first empirical characterisation. In the Week 4 interaction analysis, overlap and disagreement co-occurred in 100.0% of the common 20-profile subset. Raw disagreement resolution produced an average of 3.45 duplicate or recheck invocations. After applying LCE and Theory-of-Mind role prompts, the remaining invocations fell to 1.00. These results suggest that overlap control is not only an efficiency improvement; it also reduces the amount of renewed interaction created after disagreement resolution.

For HDFC Bank's credit-risk context, this matters because multi-agent AI systems should not be evaluated only by final approval accuracy. A system may appear accurate while still wasting computation, repeatedly checking the same evidence, or pushing agents toward unstable decision behaviour. The cross-phenomenon layer therefore provides a more operationally realistic view: it asks whether the pipeline is stable, non-redundant, and explainable across stages. This is especially important in lending, where final decisions must be defensible and where inconsistent treatment of borrowers can create model-risk and governance concerns."""


def setup_doc_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 117, 182)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(90, 90, 90)


def add_body_text(doc: Document, text: str) -> None:
    for paragraph_text in text.split("\n\n"):
        paragraph = doc.add_paragraph(paragraph_text)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.08


def word_count(text: str) -> int:
    return len([part for part in text.replace("->", " ").split() if part.strip()])


def write_docx(path: Path, title: str, text: str, target_note: str) -> None:
    doc = Document()
    setup_doc_style(doc)
    add_title(
        doc,
        title,
        "HDFC Bank Internship Project | Multi-Agent Credit Risk Simulation | Week 5 Monday",
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(f"{target_note} | Word count: {word_count(text)}")
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()
    add_body_text(doc, text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_markdown(path: Path, title: str, text: str) -> None:
    path.write_text(f"# {title}\n\n{text}\n", encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    write_docx(
        SECTION_1_DOCX,
        SECTION_1_TITLE,
        SECTION_1_TEXT,
        "Required target: 300+ words",
    )
    write_markdown(SECTION_1_MD, SECTION_1_TITLE, SECTION_1_TEXT)

    write_docx(
        SECTION_24_DOCX,
        SECTION_24_TITLE,
        SECTION_24_TEXT,
        "Required target: 400+ words",
    )
    write_markdown(SECTION_24_MD, SECTION_24_TITLE, SECTION_24_TEXT)

    print("Generated Week 5 Monday writing sections:")
    print(f"{SECTION_1_DOCX} ({word_count(SECTION_1_TEXT)} words)")
    print(f"{SECTION_24_DOCX} ({word_count(SECTION_24_TEXT)} words)")
    print(f"{SECTION_1_MD}")
    print(f"{SECTION_24_MD}")


if __name__ == "__main__":
    main()
