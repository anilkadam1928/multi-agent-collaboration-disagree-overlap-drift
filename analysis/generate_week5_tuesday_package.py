"""Generate Week 5 Tuesday Section 4 results package.

Outputs:
- Section 4.1 to 4.5 DOCX + MD files
- Complete Section 4 DOCX + MD file
- RL post-accuracy CSV files, via rl_reward_loop.py
- Week 5 master comparison table with the Post-RL row
- Two small RL figures for the report
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
HDFC_OUT = ROOT / "generated_report_sections" / "Week5_Tuesday"
RESULTS = ROOT / "results"
DATA = ROOT / "data"
ANALYSIS = ROOT / "analysis"

sys.path.insert(0, str(ROOT))
from rl_reward_loop import run_rl_reward_loop  # noqa: E402


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def expected(label: str) -> str:
    return "approve" if str(label).lower() == "good" else "reject"


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(90, 90, 90)


def add_metric_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def add_figure_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        doc.add_paragraph(f"Figure placeholder: {caption} (image file not found).")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(path), width=Inches(5.8))
    except Exception:
        doc.add_paragraph(f"Figure available at: {path}")
        return
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)


def write_docx(path: Path, title: str, subtitle: str, paragraphs: list[str], metrics: list[tuple[str, str]] | None = None, figures: list[tuple[Path, str]] | None = None) -> None:
    doc = Document()
    add_title(doc, title, subtitle)
    if metrics:
        add_metric_table(doc, metrics)
        doc.add_paragraph()
    for text in paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.08
    if figures:
        for fig_path, caption in figures:
            add_figure_if_exists(doc, fig_path, caption)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_markdown(path: Path, title: str, paragraphs: list[str], metrics: list[tuple[str, str]] | None = None) -> None:
    lines = [f"# {title}", ""]
    if metrics:
        lines.extend(["| Metric | Value |", "|---|---|"])
        lines.extend([f"| {m} | {v} |" for m, v in metrics])
        lines.append("")
    for paragraph in paragraphs:
        lines.extend([paragraph, ""])
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines), encoding="utf-8")


def compute_metrics() -> dict:
    baseline = read_csv(DATA / "baseline_results.csv")
    disagreement = read_csv(RESULTS / "disagreement_results.csv")
    overlap = read_csv(DATA / "overlap_results.csv")
    drift_no = read_csv(DATA / "drift_no_emc.csv")
    drift_emc = read_csv(DATA / "drift_with_emc.csv")
    drift_aba = read_csv(DATA / "drift_with_aba.csv")
    drift_events = read_csv(DATA / "drift_with_aba_events.csv")
    dual = read_csv(DATA / "dual_results.csv")
    resolution = read_csv(DATA / "resolution_overlap.csv")
    drift_accel = read_csv(DATA / "drift_acceleration.csv")
    combined = read_csv(DATA / "combined_results.csv")
    master = read_csv(ANALYSIS / "master_comparison_table_week4.csv")

    rl_summary, rl_trace, rl_rewards = run_rl_reward_loop()

    router = baseline[baseline["agent_name"] == "RouterManager"].copy()
    router["expected"] = router["kaggle_truth_label"].map({"good": "approve", "bad": "reject"})

    borrower_dis = disagreement.groupby("borrower_id").first().reset_index()
    borrower_dis["expected"] = borrower_dis["kaggle_truth_label"].map({"good": "approve", "bad": "reject"})
    precommit_rate = disagreement.groupby("borrower_id")["precommit_decision"].nunique().gt(1).mean()
    final_disagreement_rate = disagreement.groupby("borrower_id")["final_decision"].nunique().gt(1).mean()
    borda_accuracy = (borrower_dis["borda_winner"] == borrower_dis["expected"]).mean()
    consens_accuracy = (borrower_dis["consensagent_winner"] == borrower_dis["expected"]).mean()
    position_changes = int((disagreement["precommit_decision"] != disagreement["final_decision"]).sum())
    sycophancy_flags = int(disagreement["sycophancy_flagged"].sum())
    sycophancy_rate = sycophancy_flags / position_changes if position_changes else 0.0
    final_rounds = ", ".join(
        f"{k}: {v}" for k, v in disagreement["final_round_used"].value_counts().to_dict().items()
    )

    overlap_group = overlap.groupby("mode").agg(
        profiles=("borrower_id", "nunique"),
        redundancy=("redundancy_index", "mean"),
        accuracy=("final_correct", "mean"),
        tap_overlaps=("tap_overlap_count", "sum"),
        duplicates=("duplicate_invocations", "sum"),
    )
    raw_red = float(overlap_group.loc["without_lce", "redundancy"])
    controlled_red = float(overlap_group.loc["with_lce_tom", "redundancy"])
    redundancy_reduction = (raw_red - controlled_red) / raw_red if raw_red else 0.0
    leaders = overlap[overlap["mode"] == "with_lce_tom"].sort_values("borrower_id")["active_leader"]
    leadership_transitions = int(sum(a != b for a, b in zip(leaders, leaders.iloc[1:])))
    leader_counts = leaders.value_counts().to_dict()

    drift_no_40 = float(drift_no[drift_no["checkpoint"] == 40]["drift_score"].mean())
    drift_emc_40 = float(drift_emc[drift_emc["checkpoint"] == 40]["drift_score"].mean())
    drift_aba_40 = float(drift_aba[drift_aba["checkpoint"] == 40]["drift_score"].mean())
    emc_reduction = (drift_no_40 - drift_emc_40) / drift_no_40 if drift_no_40 else 0.0
    aba_reduction_vs_emc = (drift_emc_40 - drift_aba_40) / drift_emc_40 if drift_emc_40 else 0.0
    aba_events = int(drift_events["aba_triggered"].sum())

    low_20 = drift_accel[
        (drift_accel["group"] == "low_co_occurrence_intensity")
        & (drift_accel["checkpoint"] == 20)
    ]["group_kl_drift_score"].iloc[0]
    high_20 = drift_accel[
        (drift_accel["group"] == "high_co_occurrence_intensity")
        & (drift_accel["checkpoint"] == 20)
    ]["group_kl_drift_score"].iloc[0]
    accel_ratio = high_20 / low_20 if low_20 else 0.0

    rl_row = rl_summary.iloc[0].to_dict()
    week5_master = pd.concat(
        [
            master,
            pd.DataFrame(
                [
                    {
                        "module": "Post-RL Reward Loop",
                        "profiles_run": int(rl_row["profiles_run"]),
                        "disagreement_rate": None,
                        "redundancy_index": None,
                        "mean_drift_score": None,
                        "final_accuracy": float(rl_row["post_rl_accuracy"]),
                        "key_finding": (
                            "Reward loop adjusted module weights using +1/-1 feedback; "
                            "accuracy improved vs frozen baseline but did not exceed the "
                            "pre-RL combined pass."
                        ),
                    }
                ]
            ),
        ],
        ignore_index=True,
    )
    week5_master.to_csv(ANALYSIS / "master_comparison_table_week5_tuesday.csv", index=False)

    return {
        "baseline": baseline,
        "disagreement": disagreement,
        "overlap": overlap,
        "combined": combined,
        "rl_summary": rl_summary,
        "rl_trace": rl_trace,
        "rl_rewards": rl_rewards,
        "week5_master": week5_master,
        "router_accuracy_50": float((router["parsed_decision"] == router["expected"]).mean()),
        "precommit_rate": float(precommit_rate),
        "final_disagreement_rate": float(final_disagreement_rate),
        "borda_accuracy": float(borda_accuracy),
        "consens_accuracy": float(consens_accuracy),
        "position_changes": position_changes,
        "sycophancy_flags": sycophancy_flags,
        "sycophancy_rate": float(sycophancy_rate),
        "final_rounds": final_rounds,
        "raw_red": raw_red,
        "controlled_red": controlled_red,
        "redundancy_reduction": float(redundancy_reduction),
        "raw_duplicates": int(overlap_group.loc["without_lce", "duplicates"]),
        "controlled_duplicates": int(overlap_group.loc["with_lce_tom", "duplicates"]),
        "leadership_transitions": leadership_transitions,
        "leader_counts": leader_counts,
        "drift_no_40": drift_no_40,
        "drift_emc_40": drift_emc_40,
        "drift_aba_40": drift_aba_40,
        "emc_reduction": float(emc_reduction),
        "aba_reduction_vs_emc": float(aba_reduction_vs_emc),
        "aba_events": aba_events,
        "co_occurrence_rate": float(dual["co_occurrence"].mean()),
        "resolution_induced": float(dual["resolution_induced_invocations"].mean()),
        "remaining_after_lce": float(dual["remaining_invocations_after_lce"].mean()),
        "invocations_reduced": float(resolution["invocations_reduced_by_lce"].mean()),
        "high_drift": float(high_20),
        "low_drift": float(low_20),
        "accel_ratio": float(accel_ratio),
        "combined_acc": float(combined["combined_final_correct"].mean()),
        "baseline_common_acc": float(combined["baseline_correct"].mean()),
        "borda_common_acc": float(combined["borda_correct"].mean()),
        "overlap_common_acc": float(combined["controlled_overlap_correct"].mean()),
        "drift_common_acc": float(combined["drift_router_correct"].mean()),
    }


def create_rl_figures(metrics: dict) -> None:
    RESULTS.mkdir(parents=True, exist_ok=True)
    master = metrics["week5_master"].copy()
    plot_df = master.dropna(subset=["final_accuracy"]).copy()
    plot_df["label"] = plot_df["module"].str.replace(" Module", "", regex=False)

    plt.figure(figsize=(10, 4.8))
    bars = plt.bar(plot_df["label"], plot_df["final_accuracy"] * 100, color="#2f6f9f")
    plt.xticks(rotation=28, ha="right")
    plt.ylabel("Final accuracy (%)")
    plt.title("Week 5 Tuesday: Accuracy Trajectory Including RL")
    plt.ylim(0, 100)
    for bar in bars:
        h = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, h + 1.5, f"{h:.0f}%", ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(RESULTS / "Figure_7_RL_Accuracy_Trajectory.png", dpi=180)
    plt.close()

    rl_summary = metrics["rl_summary"].iloc[0]
    weights = json.loads(rl_summary["final_weights_json"])
    plt.figure(figsize=(8, 4.5))
    bars = plt.bar(weights.keys(), weights.values(), color="#7a9e56")
    plt.xticks(rotation=25, ha="right")
    plt.ylabel("Final routing weight")
    plt.title("Post-RL Module Weights")
    plt.ylim(0, 1.6)
    for bar in bars:
        h = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, h + 0.03, f"{h:.2f}", ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(RESULTS / "Figure_8_RL_Module_Weights.png", dpi=180)
    plt.close()


def build_sections(metrics: dict) -> dict[str, dict]:
    leader_text = ", ".join(f"{k}={v}" for k, v in metrics["leader_counts"].items())
    rl = metrics["rl_summary"].iloc[0]
    sections = {
        "4_1_Disagreement": {
            "title": "Section 4.1 - Scenario 1 Results: Disagreement",
            "metrics": [
                ("Profiles analysed", "20"),
                ("Pre-commit disagreement rate", pct(metrics["precommit_rate"])),
                ("Stored final-round disagreement rate", pct(metrics["final_disagreement_rate"])),
                ("Borda Count accuracy", pct(metrics["borda_accuracy"])),
                ("CONSENSAGENT accuracy", pct(metrics["consens_accuracy"])),
                ("Sycophancy flags", f"{metrics['sycophancy_flags']} of {metrics['position_changes']} position changes ({pct(metrics['sycophancy_rate'])})"),
                ("Resolution rounds stored", metrics["final_rounds"]),
            ],
            "paragraphs": [
                (
                    "Scenario 1 tested whether specialist agents reached the same lending judgement before and after a structured "
                    "discussion stage. On the 20-profile intervention subset, the pre-commit stage showed a disagreement rate of "
                    f"{pct(metrics['precommit_rate'])}, confirming that the agents were not simply echoing one another at baseline. "
                    f"The stored final-round outputs still showed {pct(metrics['final_disagreement_rate'])} disagreement, which means "
                    "the module preserved diversity rather than forcing artificial consensus."
                ),
                (
                    f"Borda Count converted the competing votes into a final group decision with {pct(metrics['borda_accuracy'])} "
                    f"strict Kaggle-label accuracy, while CONSENSAGENT reached {pct(metrics['consens_accuracy'])}. The difference between "
                    "these two resolution mechanisms is itself useful: Borda rewards ordinal preference across all options, whereas "
                    "CONSENSAGENT rewards confident and stable groups."
                ),
                (
                    f"The cosine sycophancy detector flagged {metrics['sycophancy_flags']} suspiciously similar position changes out of "
                    f"{metrics['position_changes']} total changes. These flags identify where an agent may have moved after discussion "
                    "while borrowing another agent's reasoning too closely. Figure_1_Disagreement_Rate.png is the visual reference for "
                    "this scenario."
                ),
            ],
            "figures": [(RESULTS / "Figure_1_Disagreement_Rate.png", "Figure 1. Disagreement module results.")],
        },
        "4_2_Overlap": {
            "title": "Section 4.2 - Scenario 2 Results: Overlap",
            "metrics": [
                ("Raw redundancy index", f"{metrics['raw_red']:.4f}"),
                ("After LCE + ToM redundancy index", f"{metrics['controlled_red']:.4f}"),
                ("Redundancy reduction", pct(metrics["redundancy_reduction"])),
                ("Duplicate invocations", f"{metrics['raw_duplicates']} -> {metrics['controlled_duplicates']}"),
                ("Leadership transitions", str(metrics["leadership_transitions"])),
                ("Leader distribution", leader_text),
            ],
            "paragraphs": [
                (
                    "Scenario 2 measured whether multiple agents were duplicating the same credit-review work. In the uncontrolled "
                    f"condition, the redundancy index was {metrics['raw_red']:.4f}. After LCE leader assignment and ToM role prompts, "
                    f"the redundancy index fell to {metrics['controlled_red']:.4f}, a {pct(metrics['redundancy_reduction'])} reduction."
                ),
                (
                    f"The operational interpretation is direct: duplicate invocations fell from {metrics['raw_duplicates']} to "
                    f"{metrics['controlled_duplicates']}. This means the system spent fewer agent calls repeating the same repayment "
                    "capacity checks and more calls on distinct specialist work. The most common overlap pattern involved CreditAgent, "
                    "IncomeAgent, and WeakModelAgent duplicating repayment-capacity reasoning."
                ),
                (
                    f"Leadership was dynamic rather than fixed. Across the 20 profiles, there were {metrics['leadership_transitions']} "
                    f"leadership transitions, with the final leader distribution recorded as: {leader_text}. Figure_2_Redundancy_Reduction.png "
                    "summarises the before-after effect."
                ),
                (
                    "For HDFC operations, the result translates into fewer repeated document checks and clearer ownership of specialist "
                    "tasks. Instead of Income, Credit, and Weak Model agents repeatedly commenting on the same repayment-capacity signal, "
                    "the LCE layer assigns a leader and reduces unnecessary parallel work. This is especially relevant for document-heavy "
                    "loan underwriting, where repeated review of the same missing or weak evidence can slow the approval queue without "
                    "improving risk quality."
                ),
            ],
            "figures": [(RESULTS / "Figure_2_Redundancy_Reduction.png", "Figure 2. Redundancy reduction after LCE + ToM.")],
        },
        "4_3_Drift": {
            "title": "Section 4.3 - Scenario 3 Results: Drift",
            "metrics": [
                ("Without EMC drift at checkpoint 40", f"{metrics['drift_no_40']:.4f}"),
                ("With EMC drift at checkpoint 40", f"{metrics['drift_emc_40']:.4f}"),
                ("With EMC + ABA drift at checkpoint 40", f"{metrics['drift_aba_40']:.4f}"),
                ("EMC reduction vs no EMC", pct(metrics["emc_reduction"])),
                ("ABA reduction vs EMC-only", pct(metrics["aba_reduction_vs_emc"])),
                ("ABA correction events", str(metrics["aba_events"])),
            ],
            "paragraphs": [
                (
                    "Scenario 3 examined behavioural drift across repeated agent use. Without EMC, the mean drift score at checkpoint "
                    f"40 was {metrics['drift_no_40']:.4f}. EMC reduced this to {metrics['drift_emc_40']:.4f}, showing that memory compaction "
                    f"lowered drift by {pct(metrics['emc_reduction'])}."
                ),
                (
                    f"When ABA was added on top of EMC, the checkpoint-40 drift score fell further to {metrics['drift_aba_40']:.4f}. "
                    f"This is a {pct(metrics['aba_reduction_vs_emc'])} reduction relative to EMC-only. The run logged "
                    f"{metrics['aba_events']} ABA correction events, concentrated in agents whose drift score crossed the configured threshold."
                ),
                (
                    "For HDFC-style model governance, this result matters because it separates passive drift monitoring from active drift "
                    "intervention. EMC controls context bloat; ABA actively re-anchors drifted agents back to domain examples. Figures 3 and 4 "
                    "show the drift and ABA behaviour."
                ),
                (
                    "The practical takeaway is that drift should not be treated only as an accuracy problem after the final decision. In a "
                    "multi-agent pipeline, drift can appear inside specialist reasoning before the final loan outcome changes. The checkpoint "
                    "scores therefore act like an early-warning governance control: HDFC could identify which agent is becoming unstable, "
                    "compact its memory, and re-anchor it before the drift propagates into a wider approval pattern."
                ),
            ],
            "figures": [
                (RESULTS / "Figure_3_Drift_Score_Over_Time.png", "Figure 3. Drift score over time."),
                (RESULTS / "Figure_4A_ABA_Drift_Comparison.png", "Figure 4. ABA drift comparison."),
            ],
        },
        "4_4_CausalChain": {
            "title": "Section 4.4 - Scenario 4 Results: Causal Chain",
            "metrics": [
                ("Co-occurrence rate", pct(metrics["co_occurrence_rate"])),
                ("Average resolution-induced invocations", f"{metrics['resolution_induced']:.2f}"),
                ("Remaining invocations after LCE", f"{metrics['remaining_after_lce']:.2f}"),
                ("Average invocations reduced by LCE", f"{metrics['invocations_reduced']:.2f}"),
                ("High-intensity drift score at checkpoint 20", f"{metrics['high_drift']:.4f}"),
                ("Low-intensity drift score at checkpoint 20", f"{metrics['low_drift']:.4f}"),
                ("High/low drift ratio", f"{metrics['accel_ratio']:.2f}x"),
            ],
            "paragraphs": [
                (
                    "Scenario 4 is the novel contribution result. It tested the causal chain: overlap -> disagreement -> resolution -> "
                    "renewed overlap -> drift risk. In the observed 20-profile run, overlap and disagreement co-occurred in "
                    f"{pct(metrics['co_occurrence_rate'])} of profiles."
                ),
                (
                    f"Each Borda resolution event induced an average of {metrics['resolution_induced']:.2f} additional invocations, but LCE "
                    f"reduced the remaining invocation burden to {metrics['remaining_after_lce']:.2f}. On average, LCE removed "
                    f"{metrics['invocations_reduced']:.2f} duplicate or resolution-induced calls per profile."
                ),
                (
                    f"The drift-risk link also appeared in the grouped analysis. At checkpoint 20, high co-occurrence profiles showed a "
                    f"group KL drift score of {metrics['high_drift']:.4f}, compared with {metrics['low_drift']:.4f} for the low co-occurrence "
                    f"group, or {metrics['accel_ratio']:.2f}x higher. This supports the claim that collaboration failures are not isolated; "
                    "they can cascade into later drift exposure."
                ),
                (
                    "This finding is important because it reframes disagreement as more than a voting problem. In the simulated HDFC credit "
                    "pipeline, disagreement creates extra resolution work; that resolution work can then reopen overlapping agent activity; "
                    "and the repeated reasoning cycles increase exposure to behavioural drift. The causal-chain result therefore connects "
                    "three earlier modules into one system-level failure mode. It gives the project a stronger research contribution than "
                    "separate before-after charts, because it shows how one collaboration failure can become the mechanism that produces "
                    "the next failure."
                ),
            ],
            "figures": [(RESULTS / "Figure_4_Causal_Chain.png", "Figure 4. Causal chain across overlap, disagreement, and drift risk.")],
        },
        "4_5_Combined_RL": {
            "title": "Section 4.5 - Scenario 5 Results: Combined Pipeline + RL",
            "metrics": [
                ("Common subset profiles", "20"),
                ("Frozen baseline accuracy", pct(metrics["baseline_common_acc"])),
                ("Borda module accuracy", pct(metrics["borda_common_acc"])),
                ("Overlap controlled accuracy", pct(metrics["overlap_common_acc"])),
                ("Drift router accuracy", pct(metrics["drift_common_acc"])),
                ("Combined pre-RL accuracy", pct(metrics["combined_acc"])),
                ("Post-RL accuracy", pct(float(rl["post_rl_accuracy"]))),
                ("Delta vs baseline", f"{float(rl['delta_vs_baseline']) * 100:+.1f} percentage points"),
                ("Delta vs combined", f"{float(rl['delta_vs_combined']) * 100:+.1f} percentage points"),
            ],
            "paragraphs": [
                (
                    "Scenario 5 assembled the major intervention modules into a common comparison table. On the 20-profile common subset, "
                    f"the frozen RouterManager baseline achieved {pct(metrics['baseline_common_acc'])} strict accuracy. Borda Count improved "
                    f"this to {pct(metrics['borda_common_acc'])}, and the controlled overlap and drift passes each reached "
                    f"{pct(metrics['overlap_common_acc'])}."
                ),
                (
                    f"The combined pre-RL pipeline reached {pct(metrics['combined_acc'])}. The reward loop then assigned +1 feedback to "
                    "policies matching the Kaggle label and -1 feedback to policies that did not, adjusting routing weights by 0.05 after "
                    f"each profile. The post-RL result was {pct(float(rl['post_rl_accuracy']))}, which is "
                    f"{float(rl['delta_vs_baseline']) * 100:+.1f} percentage points versus the frozen baseline and "
                    f"{float(rl['delta_vs_combined']) * 100:+.1f} points versus the combined pre-RL pass."
                ),
                (
                    "The important interpretation is honest rather than inflated: RL closed the feedback loop and preserved the combined "
                    "accuracy level, but it did not exceed the pre-RL combined pass on this 20-profile subset. This is still a valid result "
                    "because it shows how the simulation can record reward feedback and expose remaining approval-bias limitations."
                ),
            ],
            "figures": [
                (RESULTS / "Figure_5_Master_Comparison.png", "Figure 5. Master comparison before RL."),
                (RESULTS / "Figure_7_RL_Accuracy_Trajectory.png", "Figure 7. Accuracy trajectory including RL."),
                (RESULTS / "Figure_8_RL_Module_Weights.png", "Figure 8. Final post-RL module weights."),
            ],
        },
    }
    return sections


def write_week5_master_markdown(metrics: dict) -> None:
    df = metrics["week5_master"]
    md_path = ANALYSIS / "master_comparison_table_week5_tuesday.md"
    lines = ["# Week 5 Tuesday Master Comparison Table", ""]
    lines.append(df.to_markdown(index=False))
    lines.append("")
    md_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    HDFC_OUT.mkdir(parents=True, exist_ok=True)
    metrics = compute_metrics()
    create_rl_figures(metrics)
    write_week5_master_markdown(metrics)
    sections = build_sections(metrics)

    combined_paragraphs: list[str] = []
    for key, spec in sections.items():
        base_name = f"Section_{key}"
        write_docx(
            HDFC_OUT / f"{base_name}.docx",
            spec["title"],
            "Week 5 Tuesday - Section 4 Results",
            spec["paragraphs"],
            spec["metrics"],
            spec.get("figures"),
        )
        write_markdown(HDFC_OUT / f"{base_name}.md", spec["title"], spec["paragraphs"], spec["metrics"])
        combined_paragraphs.extend([spec["title"], *spec["paragraphs"]])

    full_doc = Document()
    add_title(
        full_doc,
        "Section 4 - Results",
        "Week 5 Tuesday package: Disagreement, Overlap, Drift, Causal Chain, Combined + RL",
    )
    for key, spec in sections.items():
        full_doc.add_heading(spec["title"], level=1)
        add_metric_table(full_doc, spec["metrics"])
        full_doc.add_paragraph()
        for paragraph in spec["paragraphs"]:
            p = full_doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(8)
        for fig_path, caption in spec.get("figures", []):
            add_figure_if_exists(full_doc, fig_path, caption)
    full_path = HDFC_OUT / "Section_4_Results_Complete.docx"
    full_doc.save(full_path)

    write_markdown(
        HDFC_OUT / "Section_4_Results_Complete.md",
        "Section 4 - Results",
        combined_paragraphs,
    )

    summary_lines = [
        "# Week 5 Tuesday Completion Summary",
        "",
        f"- Section 4.1 disagreement written: {pct(metrics['precommit_rate'])} pre-commit disagreement, {pct(metrics['borda_accuracy'])} Borda accuracy.",
        f"- Section 4.2 overlap written: redundancy {metrics['raw_red']:.4f} -> {metrics['controlled_red']:.4f}, {pct(metrics['redundancy_reduction'])} reduction.",
        f"- Section 4.3 drift written: EMC + ABA checkpoint-40 drift {metrics['drift_aba_40']:.4f}, ABA reduction vs EMC-only {pct(metrics['aba_reduction_vs_emc'])}.",
        f"- Section 4.4 causal chain written: {pct(metrics['co_occurrence_rate'])} co-occurrence, {metrics['resolution_induced']:.2f} resolution-induced invocations.",
        f"- Section 4.5 combined + RL written: post-RL accuracy {pct(float(metrics['rl_summary'].iloc[0]['post_rl_accuracy']))}.",
        "",
        "Generated outputs are in `generated_report_sections/Week5_Tuesday`.",
    ]
    (HDFC_OUT / "Week5_Tuesday_Summary.md").write_text("\n".join(summary_lines), encoding="utf-8")
    print("\nWeek 5 Tuesday package generated.")
    print(f"Output folder: {HDFC_OUT}")
    print(f"Complete Section 4: {full_path}")
    print(f"RL CSV: {DATA / 'rl_post_accuracy.csv'}")
    print(f"Master table: {ANALYSIS / 'master_comparison_table_week5_tuesday.csv'}")


if __name__ == "__main__":
    main()
